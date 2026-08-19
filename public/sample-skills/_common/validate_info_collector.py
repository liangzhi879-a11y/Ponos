"""资料收集清单审核脚本

用途：审核高新技术企业认定申报中的资料收集清单是否合规，包括：
  1. 资料收集清单完整性：检查清单文件（docx/xlsx）是否存在
  2. 清单内容校验：是否包含所有必要资料项（IP/RD/PS/人员/财务/制度等类别）
  3. 清单格式校验：检查清单格式是否规范
  4. 补充资料清单校验：是否列出了需要补充的资料
  5. 已有资料与补充资料分离校验
  6. 清单中是否排除了财务资料（财务资料由财务模块处理）

用法：
  python validate_info_collector.py --dir "资料清单目录"
  python validate_info_collector.py --dir "资料清单目录" --project-root "项目根目录"
  python validate_info_collector.py --file "资料收集清单.xlsx" --project-root "项目根目录"

输出：JSON 格式审核报告，结构如下：
  {
    "passed": bool,
    "errors": [{"file": "...", "row": 0, "field": "...", "reason": "..."}],
    "warnings": [...],
    "stats": {"total_items": N, "existing_items": N, "supplement_items": N, ...}
  }

退出码：审核通过 0，存在错误 1

依赖：openpyxl（用于读取 xlsx 清单）；若无 openpyxl 则降级为仅校验文件存在性
      python-docx（用于读取 docx 清单）；若无 python-docx 则跳过 docx 内容校验
"""

import os
import re
import sys
import json
import argparse


# ============================================================
# 常量定义
# ============================================================

# 清单文件支持的扩展名
CHECKLIST_EXTENSIONS = ('.docx', '.xlsx', '.doc', '.xls')

# 清单文件名匹配模式
CHECKLIST_NAME_PATTERNS = [
    r'资料.*清单.*\.(xlsx|docx)$',
    r'材料.*清单.*\.(xlsx|docx)$',
    r'收集.*清单.*\.(xlsx|docx)$',
    r'清单.*\.(xlsx|docx)$',
    r'checklist.*\.(xlsx|docx)$',
]

# 必要资料类别及其关键词（非财务类）
REQUIRED_CATEGORIES = {
    '知识产权(IP)': ['知识产权', 'IP', '专利', '软著', '证书'],
    '研发活动(RD)': ['研发活动', 'RD', '研发项目', '立项书', '研发报告'],
    '高新产品(PS)': ['高新产品', 'PS', '高新技术产品', '产品服务', '技术服务'],
    '科技人员': ['科技人员', '人员', '花名册', '社保', '人员清单'],
    '管理制度': ['管理制度', '研发制度', '激励制度', '辅助账', '产学研'],
    '成果转化': ['成果转化', '科技成果', '转化证明'],
}

# 财务资料关键词（应由财务模块处理，不应出现在本清单中）
FINANCIAL_KEYWORDS = [
    '审计报告', '财务报表', '资产负债表', '利润表', '现金流量表',
    '纳税申报表', '年度纳税', '研发费用加计扣除', '所得税',
    '财务审计', '财务情况', '会计报表', '营业收入', '净资产',
]

# 补充资料标记关键词（用于识别补充资料清单）
SUPPLEMENT_KEYWORDS = ['补充', '待补充', '缺失', '待收集', '需收集', '缺', '未提供', '待提供']

# 已有资料标记关键词
EXISTING_KEYWORDS = ['已有', '已提供', '已收集', '已获得', '已齐', '√', '✓']


def _load_project_index(project_root):
    """加载 project_index.json，用于获取项目上下文信息

    查找顺序：
      1. project_root/.trae/project_knowledge/project_index.json
      2. project_root/.trae/project_index.json
    """
    if not project_root:
        return None
    candidates = [
        os.path.join(project_root, '.trae', 'project_knowledge', 'project_index.json'),
        os.path.join(project_root, '.trae', 'project_index.json'),
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except (json.JSONDecodeError, OSError):
                return None
    return None


def _find_checklist_files(checklist_dir):
    """在目录中查找清单文件

    返回：[{'file': 文件名, 'path': 完整路径, 'ext': 扩展名, 'is_supplement': 是否补充清单}, ...]
    """
    files = []
    if not os.path.isdir(checklist_dir):
        return files

    for entry in sorted(os.listdir(checklist_dir)):
        ext = os.path.splitext(entry)[1].lower()
        if ext not in CHECKLIST_EXTENSIONS:
            continue
        is_supplement = any(kw in entry for kw in SUPPLEMENT_KEYWORDS)
        files.append({
            'file': entry,
            'path': os.path.join(checklist_dir, entry),
            'ext': ext,
            'is_supplement': is_supplement,
        })
    return files


def _read_xlsx_checklist(file_path):
    """读取 xlsx 格式清单文件

    返回：(rows, error_or_None)
      rows: [{'row': 行号, 'cells': [单元格值列表], 'raw': 原始行文本}, ...]
    """
    try:
        import openpyxl
    except ImportError:
        return None, 'openpyxl 未安装，跳过 xlsx 清单读取'

    if not os.path.exists(file_path):
        return None, f'文件不存在: {file_path}'

    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
    except Exception as e:
        return None, f'加载 Excel 失败: {e}'

    ws = wb.active
    rows = []

    # 自动检测表头行
    header_row = 1
    for r in range(1, min(5, ws.max_row + 1)):
        v = ws.cell(r, 1).value
        if v and ('序号' in str(v) or '类别' in str(v) or '资料' in str(v) or '名称' in str(v)):
            header_row = r
            break

    for r in range(header_row + 1, ws.max_row + 1):
        first_cell = ws.cell(r, 1).value
        if first_cell is None or str(first_cell).strip() == '':
            continue
        # 跳过统计/说明行
        if isinstance(first_cell, str) and ('总计' in first_cell or '合计' in first_cell or '说明' in first_cell or '备注' in first_cell):
            continue

        cells = []
        for c in range(1, ws.max_column + 1):
            v = ws.cell(r, c).value
            cells.append(str(v) if v is not None else '')

        raw_text = ' '.join(cells)
        rows.append({
            'row': r,
            'cells': cells,
            'raw': raw_text,
        })

    return rows, None


def _read_docx_checklist(file_path):
    """读取 docx 格式清单文件

    返回：(lines, error_or_None)
      lines: [{'row': 段落序号, 'text': 段落文本}, ...]
    """
    try:
        import docx
    except ImportError:
        return None, 'python-docx 未安装，跳过 docx 清单读取'

    if not os.path.exists(file_path):
        return None, f'文件不存在: {file_path}'

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        return None, f'读取 docx 文件失败: {e}'

    lines = []
    # 读取表格内容
    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            raw_text = ' '.join(cells)
            if raw_text.strip():
                lines.append({
                    'row': r_idx + 1,
                    'text': raw_text,
                    'cells': cells,
                })

    # 读取段落内容
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            lines.append({
                'row': i + 1,
                'text': text,
                'cells': [text],
            })

    return lines, None


def _classify_row(raw_text, cells):
    """分类一行清单内容

    返回：{
      'category': 资料类别或 None,
      'is_supplement': 是否补充资料,
      'is_existing': 是否已有资料,
      'is_financial': 是否财务资料,
    }
    """
    result = {
        'category': None,
        'is_supplement': False,
        'is_existing': False,
        'is_financial': False,
    }

    # 类别匹配
    for category, keywords in REQUIRED_CATEGORIES.items():
        for kw in keywords:
            if kw in raw_text:
                result['category'] = category
                break
        if result['category']:
            break

    # 财务资料匹配
    for kw in FINANCIAL_KEYWORDS:
        if kw in raw_text:
            result['is_financial'] = True
            break

    # 补充资料匹配
    for kw in SUPPLEMENT_KEYWORDS:
        if kw in raw_text:
            result['is_supplement'] = True
            break

    # 已有资料匹配
    for kw in EXISTING_KEYWORDS:
        if kw in raw_text:
            result['is_existing'] = True
            break

    return result


def validate_info_collector_directory(checklist_dir, project_root=None):
    """审核资料收集清单目录

    参数：
      checklist_dir: 资料清单目录路径
      project_root: 项目根目录（用于读取 project_index.json）

    返回：审核报告 dict
    """
    errors = []
    warnings = []
    stats = {
        'total_items': 0,
        'existing_items': 0,
        'supplement_items': 0,
        'financial_items': 0,
        'found_categories': [],
        'missing_categories': [],
        'has_checklist': False,
        'has_supplement_section': False,
        'has_existing_section': False,
        'checklist_files': [],
        'content_check_enabled': False,
    }

    # ============================================================
    # 校验 1：目录结构校验 & 清单文件存在性
    # ============================================================
    if not os.path.exists(checklist_dir):
        errors.append({
            'file': checklist_dir, 'row': 0, 'field': 'directory',
            'reason': f'资料清单目录不存在: {checklist_dir}',
        })
        return {'passed': False, 'errors': errors, 'warnings': warnings, 'stats': stats}

    if not os.path.isdir(checklist_dir):
        errors.append({
            'file': checklist_dir, 'row': 0, 'field': 'directory',
            'reason': f'资料清单路径不是目录: {checklist_dir}',
        })
        return {'passed': False, 'errors': errors, 'warnings': warnings, 'stats': stats}

    # 查找清单文件
    files = _find_checklist_files(checklist_dir)
    stats['checklist_files'] = [f['file'] for f in files]

    if not files:
        errors.append({
            'file': '', 'row': 0, 'field': 'checklist',
            'reason': '资料清单目录下未找到任何清单文件（docx/xlsx）',
        })
        return {'passed': False, 'errors': errors, 'warnings': warnings, 'stats': stats}

    stats['has_checklist'] = True

    # 检测依赖库
    try:
        import openpyxl  # noqa: F401
        import docx  # noqa: F401
        stats['content_check_enabled'] = True
    except ImportError:
        stats['content_check_enabled'] = False
        warnings.append({
            'file': '', 'row': 0, 'field': 'dependency',
            'reason': 'openpyxl 或 python-docx 未安装，降级为仅校验文件存在性（内容校验被跳过）',
        })

    # ============================================================
    # 逐个校验清单文件内容
    # ============================================================
    found_categories = set()
    all_items = []          # 所有清单项
    existing_items = []     # 已有资料项
    supplement_items = []   # 补充资料项
    financial_items = []    # 财务资料项

    for f in files:
        fname = f['file']
        ext = f['ext']

        # 校验 3：清单格式校验
        if ext in ('.doc', '.xls'):
            warnings.append({
                'file': fname, 'row': 0, 'field': 'format',
                'reason': f'清单使用 {ext} 格式，建议统一转换为 {"docx" if ext == ".doc" else "xlsx"} 格式',
            })

        if not stats['content_check_enabled']:
            continue

        # 读取清单内容
        rows = None
        if ext in ('.xlsx', '.xls'):
            rows, err = _read_xlsx_checklist(f['path'])
            if err:
                warnings.append({
                    'file': fname, 'row': 0, 'field': 'content',
                    'reason': err,
                })
                continue
            # xlsx 行数据
            for row_data in rows:
                raw_text = row_data['raw']
                if not raw_text.strip():
                    continue
                classified = _classify_row(raw_text, row_data['cells'])
                item = {
                    'file': fname,
                    'row': row_data['row'],
                    'text': raw_text,
                    **classified,
                }
                all_items.append(item)

        elif ext in ('.docx', '.doc'):
            lines, err = _read_docx_checklist(f['path'])
            if err:
                warnings.append({
                    'file': fname, 'row': 0, 'field': 'content',
                    'reason': err,
                })
                continue
            for line_data in lines:
                raw_text = line_data['text']
                if not raw_text.strip():
                    continue
                classified = _classify_row(raw_text, line_data.get('cells', [raw_text]))
                item = {
                    'file': fname,
                    'row': line_data['row'],
                    'text': raw_text,
                    **classified,
                }
                all_items.append(item)

    # 统计分类结果
    for item in all_items:
        if item['category']:
            found_categories.add(item['category'])
        if item['is_supplement']:
            supplement_items.append(item)
        if item['is_existing']:
            existing_items.append(item)
        if item['is_financial']:
            financial_items.append(item)

    stats['total_items'] = len(all_items)
    stats['existing_items'] = len(existing_items)
    stats['supplement_items'] = len(supplement_items)
    stats['financial_items'] = len(financial_items)
    stats['found_categories'] = sorted(found_categories)

    # ============================================================
    # 校验 2：清单内容校验 - 必要资料类别完整性
    # ============================================================
    missing_categories = [cat for cat in REQUIRED_CATEGORIES if cat not in found_categories]
    stats['missing_categories'] = missing_categories

    for cat in missing_categories:
        errors.append({
            'file': '', 'row': 0, 'field': 'category',
            'reason': f'清单缺少必要资料类别: {cat}（关键词: {REQUIRED_CATEGORIES[cat][:3]}）',
        })

    # ============================================================
    # 校验 4：补充资料清单校验
    # ============================================================
    stats['has_supplement_section'] = len(supplement_items) > 0
    if not supplement_items:
        warnings.append({
            'file': '', 'row': 0, 'field': 'supplement',
            'reason': '清单中未识别到补充资料项（建议明确标注需要补充的资料）',
        })

    # ============================================================
    # 校验 5：已有资料与补充资料分离校验
    # ============================================================
    stats['has_existing_section'] = len(existing_items) > 0
    if existing_items and supplement_items:
        # 检查是否有行同时标记为已有和补充（矛盾）
        for item in all_items:
            if item['is_supplement'] and item['is_existing']:
                warnings.append({
                    'file': item['file'], 'row': item['row'], 'field': 'conflict',
                    'reason': f'第{item["row"]}行 同时标记为"已有"和"补充"，存在矛盾',
                })
    elif not existing_items and supplement_items:
        warnings.append({
            'file': '', 'row': 0, 'field': 'separation',
            'reason': '清单中仅有补充资料，未标记已有资料，建议明确区分已有和待补充',
        })

    # 检查是否有独立的补充资料清单文件
    has_separate_supplement_file = any(f['is_supplement'] for f in files)
    if not has_separate_supplement_file and supplement_items:
        # 在同一文件中区分了已有和补充
        pass
    elif has_separate_supplement_file:
        stats['has_supplement_section'] = True

    # ============================================================
    # 校验 6：清单中是否排除了财务资料
    # ============================================================
    if financial_items:
        for item in financial_items:
            errors.append({
                'file': item['file'], 'row': item['row'], 'field': 'financial',
                'reason': f'第{item["row"]}行 包含财务资料: "{item["text"][:50]}"（财务资料应由财务模块处理，不应在本清单中）',
            })

    # 补充资料类别覆盖校验
    if supplement_items:
        supplement_categories = set()
        for item in supplement_items:
            if item['category']:
                supplement_categories.add(item['category'])
        # 检查补充资料是否覆盖了缺失的类别
        for cat in missing_categories:
            if cat not in supplement_categories:
                warnings.append({
                    'file': '', 'row': 0, 'field': 'supplement_coverage',
                    'reason': f'缺少的类别"{cat}"未在补充资料清单中列出',
                })

    passed = len(errors) == 0
    return {
        'passed': passed,
        'errors': errors,
        'warnings': warnings,
        'stats': stats,
    }


def validate_single_info_collector_file(file_path, project_root=None):
    """审核单个资料收集清单文件"""
    errors = []
    warnings = []
    stats = {
        'total_items': 0,
        'existing_items': 0,
        'supplement_items': 0,
        'financial_items': 0,
        'found_categories': [],
        'missing_categories': [],
        'has_checklist': True,
        'has_supplement_section': False,
        'has_existing_section': False,
        'checklist_files': [],
        'content_check_enabled': False,
    }

    if not os.path.exists(file_path):
        errors.append({
            'file': file_path, 'row': 0, 'field': 'file',
            'reason': f'文件不存在: {file_path}',
        })
        return {'passed': False, 'errors': errors, 'warnings': warnings, 'stats': stats}

    fname = os.path.basename(file_path)
    ext = os.path.splitext(fname)[1].lower()
    stats['checklist_files'] = [fname]

    if ext not in CHECKLIST_EXTENSIONS:
        errors.append({
            'file': fname, 'row': 0, 'field': 'format',
            'reason': f'不支持的清单格式: {ext}（仅支持 docx/xlsx）',
        })
        return {'passed': False, 'errors': errors, 'warnings': warnings, 'stats': stats}

    # 格式校验
    if ext in ('.doc', '.xls'):
        warnings.append({
            'file': fname, 'row': 0, 'field': 'format',
            'reason': f'清单使用 {ext} 格式，建议统一转换为 {"docx" if ext == ".doc" else "xlsx"} 格式',
        })

    # 检测依赖库
    try:
        import openpyxl  # noqa: F401
        import docx  # noqa: F401
        stats['content_check_enabled'] = True
    except ImportError:
        stats['content_check_enabled'] = False
        warnings.append({
            'file': fname, 'row': 0, 'field': 'dependency',
            'reason': 'openpyxl 或 python-docx 未安装，降级为仅校验文件存在性',
        })

    if not stats['content_check_enabled']:
        passed = len(errors) == 0
        return {
            'passed': passed,
            'errors': errors,
            'warnings': warnings,
            'stats': stats,
        }

    # 读取清单内容
    all_items = []
    if ext in ('.xlsx', '.xls'):
        rows, err = _read_xlsx_checklist(file_path)
        if err:
            warnings.append({
                'file': fname, 'row': 0, 'field': 'content',
                'reason': err,
            })
        else:
            for row_data in rows or []:
                raw_text = row_data['raw']
                if not raw_text.strip():
                    continue
                classified = _classify_row(raw_text, row_data['cells'])
                all_items.append({
                    'file': fname,
                    'row': row_data['row'],
                    'text': raw_text,
                    **classified,
                })
    elif ext in ('.docx', '.doc'):
        lines, err = _read_docx_checklist(file_path)
        if err:
            warnings.append({
                'file': fname, 'row': 0, 'field': 'content',
                'reason': err,
            })
        else:
            for line_data in lines or []:
                raw_text = line_data['text']
                if not raw_text.strip():
                    continue
                classified = _classify_row(raw_text, line_data.get('cells', [raw_text]))
                all_items.append({
                    'file': fname,
                    'row': line_data['row'],
                    'text': raw_text,
                    **classified,
                })

    # 统计
    found_categories = set()
    existing_items = []
    supplement_items = []
    financial_items = []

    for item in all_items:
        if item['category']:
            found_categories.add(item['category'])
        if item['is_supplement']:
            supplement_items.append(item)
        if item['is_existing']:
            existing_items.append(item)
        if item['is_financial']:
            financial_items.append(item)

    stats['total_items'] = len(all_items)
    stats['existing_items'] = len(existing_items)
    stats['supplement_items'] = len(supplement_items)
    stats['financial_items'] = len(financial_items)
    stats['found_categories'] = sorted(found_categories)
    stats['has_supplement_section'] = len(supplement_items) > 0
    stats['has_existing_section'] = len(existing_items) > 0

    # 必要类别校验
    missing_categories = [cat for cat in REQUIRED_CATEGORIES if cat not in found_categories]
    stats['missing_categories'] = missing_categories
    for cat in missing_categories:
        errors.append({
            'file': fname, 'row': 0, 'field': 'category',
            'reason': f'清单缺少必要资料类别: {cat}',
        })

    # 补充资料校验
    if not supplement_items:
        warnings.append({
            'file': fname, 'row': 0, 'field': 'supplement',
            'reason': '清单中未识别到补充资料项',
        })

    # 已有/补充分离校验
    for item in all_items:
        if item['is_supplement'] and item['is_existing']:
            warnings.append({
                'file': fname, 'row': item['row'], 'field': 'conflict',
                'reason': f'第{item["row"]}行 同时标记为"已有"和"补充"，存在矛盾',
            })

    # 财务资料排除校验
    for item in financial_items:
        errors.append({
            'file': fname, 'row': item['row'], 'field': 'financial',
            'reason': f'第{item["row"]}行 包含财务资料（应由财务模块处理）',
        })

    passed = len(errors) == 0
    return {
        'passed': passed,
        'errors': errors,
        'warnings': warnings,
        'stats': stats,
    }


def main():
    """CLI 入口：解析参数并执行资料收集清单审核"""
    parser = argparse.ArgumentParser(
        description='资料收集清单审核脚本 - 校验高企认定资料收集清单合规性',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''示例：
  python validate_info_collector.py --dir "资料清单目录"
  python validate_info_collector.py --dir "资料清单目录" --project-root "项目根目录"
  python validate_info_collector.py --file "资料收集清单.xlsx" --project-root "项目根目录"
''',
    )
    parser.add_argument('--dir', help='资料收集清单所在目录')
    parser.add_argument('--file', help='单个清单文件路径（docx/xlsx）')
    parser.add_argument('--project-root', dest='project_root',
                        help='项目根目录（用于读取 project_index.json 获取项目上下文）')
    args = parser.parse_args()

    if not args.dir and not args.file:
        parser.print_help()
        sys.exit(1)

    if args.dir:
        report = validate_info_collector_directory(args.dir, args.project_root)
    else:
        report = validate_single_info_collector_file(args.file, args.project_root)

    # 输出 JSON 格式审核报告
    print(json.dumps(report, ensure_ascii=False, indent=2, default=str))
    sys.exit(0 if report['passed'] else 1)


if __name__ == '__main__':
    main()
