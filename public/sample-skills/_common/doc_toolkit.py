#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
doc_toolkit.py - 统一文档处理工具包

核心价值：agent 处理任何文档时，先调用 info 获取推荐方案，再调用 read/write/convert
不再需要尝试不同的库和方案，节省大量探索时间。

支持的文件类型：.xlsx .xls .docx .doc .pdf .txt .json .jpg .png

用法：
  python doc_toolkit.py info --file "报告.pdf"
      → 返回该文件类型的推荐处理方案、库、命令

  python doc_toolkit.py read --file "发票.xlsx" --format xlsx
      → 读取文件，输出JSON格式内容

  python doc_toolkit.py read --file "申请书.pdf" --format pdf --mode text
      → 提取PDF文本（mode: text/table/images）

  python doc_toolkit.py write --file "输出.xlsx" --format xlsx --data data.json
      → 从JSON数据生成文档

  python doc_toolkit.py convert --input "模板.doc" --to docx
      → 格式转换（.doc→.docx, .xls→.xlsx）

  python doc_toolkit.py convert-dir --input "模板目录" --to docx
      → 批量转换目录下所有文件

设计逻辑（agent可据此自主调改）：
  info: 读取tech_stack.json → 按扩展名返回条目
  read: 根据格式调用对应库（openpyxl/docx/fitz/pdfplumber）
  write: 根据格式调用对应库生成文档
  convert: .doc→.docx用win32com；.xls→.xlsx用xlrd+openpyxl
  如果脚本执行报错，agent应阅读本文件的设计逻辑，自主编写等效Python代码实现。
"""

import argparse
import json
import os
import sys
from pathlib import Path

# 确保 _common 目录在 sys.path 中（支持直接运行和外部调用）
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

COMMON_DIR = Path(__file__).parent
TECH_STACK_FILE = COMMON_DIR / "tech_stack.json"


def load_tech_stack():
    """加载技术栈配置"""
    try:
        return json.loads(TECH_STACK_FILE.read_text(encoding='utf-8'))
    except Exception:
        return {"file_type_handlers": {}}


# ============================================================
# info 子命令：获取文件类型的推荐处理方案
# ============================================================
def cmd_info(args):
    """返回某文件类型的推荐处理方案"""
    file_path = Path(args.file) if args.file else None
    if file_path and not file_path.exists() and not args.ext:
        print(f"[info] 文件不存在: {args.file}")
        print(f"[info] 提示：使用 --ext 参数查询某扩展名的处理方案")
        return False

    ext = args.ext or (file_path.suffix.lower() if file_path else "")
    tech_stack = load_tech_stack()
    handlers = tech_stack.get("file_type_handlers", {})

    if ext not in handlers:
        print(f"[info] ⚠ 扩展名 {ext} 未在 tech_stack.json 中配置")
        print(f"[info] 已支持的类型: {', '.join(sorted(handlers.keys()))}")
        return False

    handler = handlers[ext]
    print(f"\n{'='*70}")
    print(f"  文件类型: {ext}")
    print(f"{'='*70}")
    print(f"  描述:     {handler.get('description', '')}")
    print(f"  推荐库:   {handler.get('primary_library', '')}")
    print(f"  安装:     {handler.get('library_install', '')}")
    if handler.get('fallback_library'):
        print(f"  备选库:   {handler.get('fallback_library')}")
    if handler.get('platform'):
        print(f"  平台:     {handler.get('platform')}")

    if "read" in handler:
        print(f"\n  读取方案:")
        print(f"    函数:   {handler['read'].get('function', '')}")
        print(f"    示例:   {handler['read'].get('example', '')}")
        print(f"    命令:   {handler['read'].get('toolkit_command', '')}")

    if "write" in handler:
        print(f"\n  写入方案:")
        print(f"    函数:   {handler['write'].get('function', '')}")
        print(f"    命令:   {handler['write'].get('toolkit_command', '')}")

    if handler.get("common_scenarios"):
        print(f"\n  常见场景:")
        for scenario, method in handler["common_scenarios"].items():
            print(f"    {scenario}: {method}")

    # 任务推荐
    task_recs = tech_stack.get("task_based_recommendations", {})
    relevant_tasks = []
    for task_name, task_info in task_recs.items():
        if ext in str(task_info) or handler.get('primary_library', '') in str(task_info.get('library', '')):
            relevant_tasks.append((task_name, task_info))

    if relevant_tasks:
        print(f"\n  相关任务推荐:")
        for task_name, task_info in relevant_tasks:
            print(f"    【{task_name}】")
            print(f"      脚本: {task_info.get('primary', '')}")
            print(f"      命令: {task_info.get('command', '')}")

    return True


# ============================================================
# read 子命令：读取文档内容
# ============================================================
def cmd_read(args):
    """读取文档内容，输出JSON格式"""
    file_path = Path(args.file)
    if not file_path.exists():
        print(f"[read] ERROR: 文件不存在: {args.file}")
        return False

    ext = args.format or file_path.suffix.lower()
    if ext and not ext.startswith("."):
        ext = "." + ext

    try:
        if ext == ".xlsx":
            return read_xlsx(file_path, args.sheet, args.max_rows)
        elif ext == ".xls":
            return read_xls(file_path, args.sheet, args.max_rows)
        elif ext == ".docx":
            return read_docx(file_path, args.max_chars)
        elif ext == ".doc":
            return read_doc(file_path, args.max_chars)
        elif ext == ".pdf":
            return read_pdf(file_path, args.mode, args.max_pages, args.force_ocr, args.project)
        elif ext == ".txt":
            return read_txt(file_path, args.max_chars)
        elif ext == ".json":
            return read_json(file_path)
        else:
            print(f"[read] ERROR: 不支持的格式: {ext}")
            print(f"[read] 支持: .xlsx .xls .docx .doc .pdf .txt .json")
            return False
    except Exception as e:
        print(f"[read] ERROR: 读取失败: {e}")
        return False


def read_xlsx(file_path, sheet=None, max_rows=None):
    """读取xlsx文件"""
    try:
        import openpyxl
    except ImportError:
        print("[read] ERROR: openpyxl未安装。运行: pip install openpyxl")
        return False

    wb = openpyxl.load_workbook(file_path, data_only=True)
    sheets = wb.sheetnames
    target_sheet = sheet or sheets[0]
    ws = wb[target_sheet]

    rows_data = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if max_rows and i >= max_rows:
            break
        rows_data.append([str(c) if c is not None else "" for c in row])

    result = {
        "file": str(file_path),
        "format": "xlsx",
        "sheets": sheets,
        "current_sheet": target_sheet,
        "total_rows": ws.max_row,
        "total_cols": ws.max_column,
        "rows_read": len(rows_data),
        "data": rows_data,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return True


def read_xls(file_path, sheet=None, max_rows=None):
    """读取xls文件（旧格式）"""
    try:
        import xlrd
    except ImportError:
        print("[read] ERROR: xlrd未安装。运行: pip install xlrd")
        print("[read] 提示: 建议转换为xlsx格式: python doc_toolkit.py convert --input <file> --to xlsx")
        return False

    wb = xlrd.open_workbook(file_path)
    sheets = wb.sheet_names()
    target_sheet = sheet or sheets[0]
    ws = wb.sheet_by_name(target_sheet)

    rows_data = []
    for i in range(min(ws.nrows, max_rows or ws.nrows)):
        rows_data.append([str(ws.cell_value(i, j)) for j in range(ws.ncols)])

    result = {
        "file": str(file_path),
        "format": "xls",
        "sheets": sheets,
        "current_sheet": target_sheet,
        "total_rows": ws.nrows,
        "total_cols": ws.ncols,
        "rows_read": len(rows_data),
        "data": rows_data,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return True


def read_docx(file_path, max_chars=None):
    """读取docx文件"""
    try:
        import docx
    except ImportError:
        print("[read] ERROR: python-docx未安装。运行: pip install python-docx")
        return False

    doc = docx.Document(file_path)

    paragraphs = []
    total_chars = 0
    for p in doc.paragraphs:
        text = p.text
        if max_chars and total_chars + len(text) > max_chars:
            text = text[:max_chars - total_chars]
            paragraphs.append({"text": text, "truncated": True})
            break
        paragraphs.append({"text": text, "style": p.style.name if p.style else ""})
        total_chars += len(text)

    tables = []
    for t in doc.tables:
        table_data = []
        for row in t.rows:
            table_data.append([cell.text for cell in row.cells])
        tables.append(table_data)

    result = {
        "file": str(file_path),
        "format": "docx",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "total_chars": total_chars,
        "paragraphs": paragraphs,
        "tables": tables,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return True


def read_doc(file_path, max_chars=None):
    """读取doc文件（旧格式，需win32com）"""
    try:
        import win32com.client
    except ImportError:
        print("[read] ERROR: pywin32未安装。运行: pip install pywin32")
        print("[read] 提示: 建议转换为docx格式: python doc_toolkit.py convert --input <file> --to docx")
        return False

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(file_path.resolve()))
        text = doc.Content.Text
        if max_chars:
            text = text[:max_chars]
        result = {
            "file": str(file_path),
            "format": "doc",
            "total_chars": len(doc.Content.Text),
            "text": text,
        }
        print(json.dumps(result, ensure_ascii=False, indent=2))
        doc.Close()
        return True
    finally:
        word.Quit()


def read_pdf(file_path, mode="auto", max_pages=None, force_ocr=False, project_name="default"):
    """读取PDF文件

    模式（mode）:
      - text: 仅文本层提取（PyMuPDF）
      - table: 表格提取（pdfplumber）
      - ocr: 强制OCR识别（PaddleOCR）
      - auto: 智能检测，扫描件自动降级到OCR（推荐）
    """
    if mode == "text":
        return _read_pdf_text(file_path, max_pages)
    elif mode == "table":
        return _read_pdf_table(file_path, max_pages)
    elif mode == "ocr":
        return _read_pdf_ocr(file_path, force_ocr=True, project_name=project_name)
    elif mode == "auto":
        return _read_pdf_auto(file_path, max_pages, force_ocr, project_name)
    else:
        print(f"[read] ERROR: 不支持的PDF模式: {mode}（支持: text/table/ocr/auto）")
        return False


def _read_pdf_text(file_path, max_pages=None):
    """仅文本层提取"""
    try:
        import fitz
    except ImportError:
        print("[read] ERROR: PyMuPDF未安装。运行: pip install PyMuPDF")
        return False

    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        if max_pages and i >= max_pages:
            break
        pages.append({"page": i + 1, "text": page.get_text()})

    result = {
        "file": str(file_path),
        "format": "pdf",
        "mode": "text",
        "total_pages": len(doc),
        "pages_read": len(pages),
        "pages": pages,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return True


def _read_pdf_table(file_path, max_pages=None):
    """表格提取"""
    try:
        import pdfplumber
    except ImportError:
        print("[read] ERROR: pdfplumber未安装。运行: pip install pdfplumber")
        return False

    pages = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            if max_pages and i >= max_pages:
                break
            tables = page.extract_tables()
            pages.append({"page": i + 1, "tables": tables})

        total_pages = len(pdf.pages)

    result = {
        "file": str(file_path),
        "format": "pdf",
        "mode": "table",
        "total_pages": total_pages,
        "pages_read": len(pages),
        "pages": pages,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return True


def _read_pdf_ocr(file_path, force_ocr=True, project_name="default"):
    """OCR识别模式"""
    try:
        from ocr_engine import ocr_with_table
    except ImportError:
        print("[read] ERROR: ocr_engine不可用。运行: pip install paddleocr paddlepaddle")
        return False

    result = ocr_with_table(str(file_path), project_name=project_name)
    output = {
        "file": str(file_path),
        "format": "pdf",
        "mode": "ocr",
        "ocr_engine": result.get("ocr_engine", ""),
        "is_scanned": result.get("is_scanned", False),
        "scan_detection": result.get("detection", {}),
        "cache_hit": result.get("cache_hit", False),
        "pages": result.get("pages", []),
        "tables": result.get("tables", []),
        "processing_time_seconds": result.get("processing_time_seconds", 0),
    }
    print(json.dumps(output, ensure_ascii=False, indent=2))
    return True


def _read_pdf_auto(file_path, max_pages=None, force_ocr=False, project_name="default"):
    """智能自动模式：逐页检测，扫描件页用OCR，文本页用文本提取

    v2026-07-22 修复：
      1. 串联 detect_pages_by_type 逐页检测
      2. 混合型 PDF 按页面类型分流：扫描页 → OCR，文本页 → 文本提取
      3. 纯文本 PDF → 纯文本提取
      4. 纯扫描 PDF → 强制 OCR
    """
    try:
        from ocr_engine import detect_scanned_pdf
    except ImportError:
        # OCR引擎不可用，降级为纯文本提取
        print("[read] WARN: ocr_engine不可用，降级为纯文本提取")
        return _read_pdf_text(file_path, max_pages)

    # 1. 整体检测（已 v2026-07-22 修复：串联逐页检测）
    scan_info = detect_scanned_pdf(str(file_path))
    print(f"[read] 扫描件检测: is_scanned={scan_info['is_scanned']}, reason={scan_info.get('reason', '')}")

    # 2. v2026-07-22 关键修复：检查是否为混合型 PDF
    page_analysis = scan_info.get("page_analysis")
    if page_analysis and page_analysis.get("is_mixed"):
        # 混合型：文本页+扫描页 → 必须逐页处理
        print(f"[read] ⚠️ 检测到混合型 PDF：{page_analysis['text_page_count']}文本页 + "
              f"{page_analysis['scan_page_count']}扫描页，启用逐页处理")
        return _read_pdf_mixed(file_path, page_analysis, max_pages, project_name)

    # 3. 根据检测结果选择模式
    if scan_info["is_scanned"] or force_ocr:
        print(f"[read] 触发OCR模式（{'扫描件' if scan_info['is_scanned'] else '强制OCR'}）")
        return _read_pdf_ocr(file_path, force_ocr=True, project_name=project_name)
    else:
        print("[read] 使用文本层提取（非扫描件）")
        return _read_pdf_text(file_path, max_pages)


def _read_pdf_mixed(file_path, page_analysis, max_pages=None, project_name="default"):
    """混合型 PDF 处理：逐页判断，扫描页 OCR + 文本页文本提取

    v2026-07-22 新增：解决混合型 PDF 中扫描页数据丢失问题
    """
    try:
        import fitz
    except ImportError:
        print("[read] ERROR: PyMuPDF未安装")
        return False

    try:
        from ocr_engine import OCREngine, ocr_with_table
    except ImportError:
        print("[read] ERROR: ocr_engine不可用，无法处理混合型 PDF")
        return False

    doc = fitz.open(file_path)
    pages = []
    engine = OCREngine()

    for i, page in enumerate(doc):
        if max_pages and i >= max_pages:
            break

        text = page.get_text().strip()
        char_count = len(text)
        image_count = len(page.get_images())

        # 逐页判断类型
        if char_count >= 100 or (char_count >= 10 and image_count == 0):
            page_type = "text"
            page_data = {"page": i + 1, "type": "text", "text": text}
        elif image_count > 0:
            # 扫描页 → OCR
            page_type = "scan"
            try:
                ocr_result = engine.ocr_page(str(file_path), i)
                page_data = {
                    "page": i + 1,
                    "type": "scan",
                    "ocr_text": ocr_result.get("text", ""),
                    "ocr_confidence": ocr_result.get("confidence", 0),
                }
            except Exception as e:
                page_data = {
                    "page": i + 1,
                    "type": "scan",
                    "ocr_error": str(e),
                }
        else:
            page_type = "empty"
            page_data = {"page": i + 1, "type": "empty"}

        pages.append(page_data)

    doc.close()

    text_pages = [p for p in pages if p["type"] == "text"]
    scan_pages = [p for p in pages if p["type"] == "scan"]

    print(f"[read] 混合型 PDF 处理完成：{len(text_pages)}文本页 + {len(scan_pages)}扫描页（共{len(pages)}页）")

    # 检查 OCR 结果质量
    if scan_pages:
        failed_ocr = [p for p in scan_pages if p.get("ocr_error") or not p.get("ocr_text", "").strip()]
        if failed_ocr:
            print(f"[read] ⚠️ OCR 失败 {len(failed_ocr)} 页：{[p['page'] for p in failed_ocr]}")

    result = {
        "file": str(file_path),
        "format": "pdf",
        "mode": "mixed",
        "is_mixed": True,
        "total_pages": page_analysis.get("text_page_count", 0) + page_analysis.get("scan_page_count", 0),
        "pages_read": len(pages),
        "text_page_count": len(text_pages),
        "scan_page_count": len(scan_pages),
        "pages": pages,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return True


def detect_scan(file_path):
    """检测PDF是否为扫描件（含混合型PDF逐页分析）"""
    try:
        from ocr_engine import detect_scanned_pdf, OCREngine
    except ImportError:
        print("[detect-scan] ERROR: ocr_engine不可用")
        return False

    # 整体检测
    info = detect_scanned_pdf(str(file_path))

    # 逐页检测（混合型PDF）
    engine = OCREngine()
    page_info = engine.detect_pages_by_type(str(file_path))

    output = {
        "file": str(file_path),
        "overall_detection": info,
        "page_analysis": {
            "total_pages": page_info["total_pages"],
            "text_page_count": page_info["text_page_count"],
            "scan_page_count": page_info["scan_page_count"],
            "empty_page_count": page_info["empty_page_count"],
            "is_mixed": page_info["is_mixed"],
        },
        "recommendation": "mixed" if page_info["is_mixed"] else ("ocr" if info.get("is_scanned") else "text"),
    }

    print(json.dumps(output, ensure_ascii=False, indent=2))
    return True



def read_txt(file_path, max_chars=None):
    """读取txt文件"""
    text = file_path.read_text(encoding='utf-8')
    if max_chars:
        text = text[:max_chars]
    result = {"file": str(file_path), "format": "txt", "total_chars": len(text), "text": text}
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return True


def read_json(file_path):
    """读取json文件"""
    data = json.loads(file_path.read_text(encoding='utf-8'))
    result = {"file": str(file_path), "format": "json", "data": data}
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return True


# ============================================================
# write 子命令：写入文档
# ============================================================
def cmd_write(args):
    """从JSON数据生成文档"""
    file_path = Path(args.file)
    ext = args.format or file_path.suffix.lower()

    if not args.data and not args.data_file:
        print("[write] ERROR: 需要 --data 或 --data-file 参数")
        return False

    if args.data_file:
        data = json.loads(Path(args.data_file).read_text(encoding='utf-8'))
    else:
        data = json.loads(args.data)

    try:
        if ext == ".xlsx":
            return write_xlsx(file_path, data)
        elif ext == ".docx":
            return write_docx(file_path, data)
        elif ext == ".txt":
            return write_txt(file_path, data)
        elif ext == ".json":
            return write_json(file_path, data)
        else:
            print(f"[write] ERROR: 不支持的写入格式: {ext}")
            print(f"[write] 支持: .xlsx .docx .txt .json")
            return False
    except Exception as e:
        print(f"[write] ERROR: 写入失败: {e}")
        return False


def write_xlsx(file_path, data):
    """写入xlsx文件"""
    try:
        import openpyxl
    except ImportError:
        print("[write] ERROR: openpyxl未安装")
        return False

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = data.get("sheet_name", "Sheet1")

    for row in data.get("rows", []):
        ws.append(row)

    wb.save(file_path)
    print(f"[write] ✓ xlsx已写入: {file_path}")
    return True


def write_docx(file_path, data):
    """写入docx文件"""
    try:
        import docx
        from docx.shared import Pt
    except ImportError:
        print("[write] ERROR: python-docx未安装")
        return False

    doc = docx.Document()

    for item in data.get("content", []):
        if item["type"] == "heading":
            doc.add_heading(item["text"], level=item.get("level", 1))
        elif item["type"] == "paragraph":
            doc.add_paragraph(item["text"])
        elif item["type"] == "table":
            table = doc.add_table(rows=len(item["rows"]), cols=len(item["rows"][0]))
            for i, row in enumerate(item["rows"]):
                for j, cell in enumerate(row):
                    table.rows[i].cells[j].text = str(cell)

    doc.save(file_path)
    print(f"[write] ✓ docx已写入: {file_path}")
    return True


def write_txt(file_path, data):
    """写入txt文件"""
    text = data.get("text", json.dumps(data, ensure_ascii=False, indent=2))
    file_path.write_text(text, encoding='utf-8')
    print(f"[write] ✓ txt已写入: {file_path}")
    return True


def write_json(file_path, data):
    """写入json文件"""
    file_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"[write] ✓ json已写入: {file_path}")
    return True


# ============================================================
# convert 子命令：格式转换
# ============================================================
def cmd_convert(args):
    """格式转换"""
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[convert] ERROR: 输入文件不存在: {args.input}")
        return False

    src_ext = input_path.suffix.lower()
    dst_ext = "." + args.to.lstrip(".")

    if src_ext == ".doc" and dst_ext == ".docx":
        return convert_doc_to_docx(input_path)
    elif src_ext == ".xls" and dst_ext == ".xlsx":
        return convert_xls_to_xlsx(input_path)
    elif src_ext == dst_ext:
        print(f"[convert] 源格式和目标格式相同: {src_ext}")
        return True
    else:
        print(f"[convert] ERROR: 不支持的转换: {src_ext} → {dst_ext}")
        print(f"[convert] 支持: .doc→.docx, .xls→.xlsx")
        return False


def convert_doc_to_docx(input_path):
    """doc转docx"""
    try:
        import win32com.client
    except ImportError:
        print("[convert] ERROR: pywin32未安装。运行: pip install pywin32")
        return False

    output_path = input_path.with_suffix(".docx")
    if output_path.exists():
        print(f"[convert] 目标已存在，跳过: {output_path}")
        return True

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(input_path.resolve()))
        doc.SaveAs(str(output_path.resolve()), FileFormat=16)  # 16 = wdFormatXMLDocument
        doc.Close()
        print(f"[convert] ✓ {input_path.name} → {output_path.name}")
        return True
    finally:
        word.Quit()


def convert_xls_to_xlsx(input_path):
    """xls转xlsx"""
    try:
        import xlrd
        import openpyxl
    except ImportError as e:
        print(f"[convert] ERROR: 依赖缺失: {e}")
        return False

    output_path = input_path.with_suffix(".xlsx")
    if output_path.exists():
        print(f"[convert] 目标已存在，跳过: {output_path}")
        return True

    wb_old = xlrd.open_workbook(input_path)
    wb_new = openpyxl.Workbook()

    for i, sheet_name in enumerate(wb_old.sheet_names()):
        ws_old = wb_old.sheet_by_name(sheet_name)
        ws_new = wb_new.create_sheet(sheet_name if i > 0 else "Sheet1")

        for row in range(ws_old.nrows):
            for col in range(ws_old.ncols):
                ws_new.cell(row=row + 1, column=col + 1, value=ws_old.cell_value(row, col))

    # 删除默认空sheet
    if len(wb_new.sheetnames) > 1:
        del wb_new["Sheet"]

    wb_new.save(output_path)
    print(f"[convert] ✓ {input_path.name} → {output_path.name}")
    return True


def cmd_convert_dir(args):
    """批量转换目录"""
    input_dir = Path(args.input)
    if not input_dir.is_dir():
        print(f"[convert-dir] ERROR: 不是目录: {args.input}")
        return False

    target_ext = "." + args.to.lstrip(".")
    converted = 0
    failed = 0

    for filepath in input_dir.rglob("*"):
        if filepath.suffix.lower() in (".doc", ".xls") and filepath.suffix.lower() != target_ext:
            print(f"\n[convert-dir] 处理: {filepath.name}")
            args.input = str(filepath)
            args.to = args.to
            if cmd_convert(args):
                converted += 1
            else:
                failed += 1

    print(f"\n[convert-dir] 完成: 转换 {converted} 个，失败 {failed} 个")
    return failed == 0


# ============================================================
# 命令行入口
# ============================================================
def main():
    parser = argparse.ArgumentParser(
        description='统一文档处理工具包',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例：
  # 查询文件类型的推荐处理方案
  python doc_toolkit.py info --file "报告.pdf"
  python doc_toolkit.py info --ext .xlsx

  # 读取文档
  python doc_toolkit.py read --file "发票.xlsx" --format xlsx --max-rows 100
  python doc_toolkit.py read --file "申请书.pdf" --format pdf --mode text --max-pages 5
  python doc_toolkit.py read --file "立项书.docx" --format docx --max-chars 5000

  # 写入文档
  python doc_toolkit.py write --file "输出.xlsx" --format xlsx --data-file data.json
  python doc_toolkit.py write --file "输出.docx" --format docx --data '{"content":[{"type":"paragraph","text":"hello"}]}'

  # 格式转换
  python doc_toolkit.py convert --input "模板.doc" --to docx
  python doc_toolkit.py convert --input "旧表.xls" --to xlsx
  python doc_toolkit.py convert-dir --input "模板目录" --to docx
        ''',
    )
    subparsers = parser.add_subparsers(dest="command", help="子命令")

    # info 子命令
    info_parser = subparsers.add_parser("info", help="获取文件类型推荐方案")
    info_parser.add_argument("--file", default=None, help="文件路径")
    info_parser.add_argument("--ext", default=None, help="直接指定扩展名（如 .xlsx）")

    # read 子命令
    read_parser = subparsers.add_parser("read", help="读取文档内容")
    read_parser.add_argument("--file", required=True, help="文件路径")
    read_parser.add_argument("--format", default=None, help="强制指定格式（默认从扩展名推断）")
    read_parser.add_argument("--sheet", default=None, help="Excel工作表名")
    read_parser.add_argument("--mode", default="auto", help="PDF模式: auto(智能混合)/text/table/ocr（默认auto，自动检测扫描件页并OCR）")
    read_parser.add_argument("--max-rows", type=int, default=None, help="Excel最大行数")
    read_parser.add_argument("--max-pages", type=int, default=None, help="PDF最大页数")
    read_parser.add_argument("--max-chars", type=int, default=None, help="文本最大字符数")
    read_parser.add_argument("--force-ocr", action="store_true", help="PDF强制OCR（跳过文本层）")
    read_parser.add_argument("--project", default="default", help="OCR缓存项目名（项目隔离）")

    # detect-scan 子命令
    scan_parser = subparsers.add_parser("detect-scan", help="检测PDF是否为扫描件")
    scan_parser.add_argument("--file", required=True, help="PDF文件路径")

    # write 子命令
    write_parser = subparsers.add_parser("write", help="写入文档")
    write_parser.add_argument("--file", required=True, help="输出文件路径")
    write_parser.add_argument("--format", default=None, help="强制指定格式")
    write_parser.add_argument("--data", default=None, help="JSON格式数据")
    write_parser.add_argument("--data-file", default=None, help="JSON数据文件路径")

    # convert 子命令
    conv_parser = subparsers.add_parser("convert", help="格式转换")
    conv_parser.add_argument("--input", required=True, help="输入文件路径")
    conv_parser.add_argument("--to", required=True, help="目标格式（如 docx/xlsx）")

    # convert-dir 子命令
    conv_dir_parser = subparsers.add_parser("convert-dir", help="批量转换目录")
    conv_dir_parser.add_argument("--input", required=True, help="输入目录")
    conv_dir_parser.add_argument("--to", required=True, help="目标格式")

    args = parser.parse_args()

    if args.command == "info":
        success = cmd_info(args)
        sys.exit(0 if success else 1)
    elif args.command == "read":
        success = cmd_read(args)
        sys.exit(0 if success else 1)
    elif args.command == "detect-scan":
        success = detect_scan(args.file)
        sys.exit(0 if success else 1)
    elif args.command == "write":
        success = cmd_write(args)
        sys.exit(0 if success else 1)
    elif args.command == "convert":
        success = cmd_convert(args)
        sys.exit(0 if success else 1)
    elif args.command == "convert-dir":
        success = cmd_convert_dir(args)
        sys.exit(0 if success else 1)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
