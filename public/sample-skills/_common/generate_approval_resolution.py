#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_approval_resolution.py - 立项决议汇总文档生成脚本

生成一份独立的《研发项目立项决议》Word 文档，包含汇总表格 + 逐项决议正文。

用法:
  python generate_approval_resolution.py \
    --rd-data "path/to/rd_data.json" \
    --enterprise "企业名称" \
    --year 2024 \
    --output "path/to/立项决议汇总.docx"
"""

import argparse
import json
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr, val in edge_data.items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_normal_paragraph(doc, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent and alignment == WD_ALIGN_PARAGRAPH.LEFT:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def build_summary_table(doc, rd_list):
    """构建 RD 项目汇总表格"""
    headers = ['序号', '项目编号', '项目名称', '起止时间', '决议日期', '审批结果']
    col_widths = [Cm(1.2), Cm(2), Cm(5.5), Cm(3), Cm(2.5), Cm(2)]

    table = doc.add_table(rows=len(rd_list) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for c, header in enumerate(headers):
        cell = table.cell(0, c)
        cell.width = col_widths[c]
        set_cell_text(cell, header, bold=True, size=10)

    for i, rd in enumerate(rd_list):
        row = table.rows[i + 1]
        values = [
            str(i + 1),
            rd.get('code', f'RD{i+1:02d}'),
            rd.get('name', ''),
            rd.get('period', ''),
            rd.get('approval_date', datetime.now().strftime('%Y.%m')),
            '通过'
        ]
        for c, val in enumerate(values):
            cell = row.cells[c]
            cell.width = col_widths[c]
            alignment = WD_ALIGN_PARAGRAPH.LEFT if c in (1, 2, 3) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cell, val, size=10, alignment=alignment)

    return table


def generate_approval_resolution(rd_data, enterprise, year, output_path):
    """
    生成立项决议汇总文档

    rd_data: list[dict], 每个 dict 包含 code/name/period/resolution 字段
    enterprise: 企业名称
    year: 申报年度
    output_path: 输出 .docx 路径
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    font.size = Pt(12)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(18)
    run = title_p.add_run('研发项目立项决议汇总表')
    set_run_font(run, size=22, bold=True)

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(6)
    run = info_p.add_run(f'企业名称：{enterprise}    申报年度：{year}年度    科技人员数：____人')
    set_run_font(run, size=11)

    doc.add_paragraph()

    t1_p = doc.add_paragraph()
    t1_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t1_p.paragraph_format.space_before = Pt(6)
    t1_p.paragraph_format.space_after = Pt(6)
    run = t1_p.add_run(f'经公司研发部门评审，{year}年度共立项研发项目{len(rd_data)}项，具体如下：')
    set_run_font(run, size=12)

    build_summary_table(doc, rd_data)

    doc.add_paragraph()

    section_p = doc.add_paragraph()
    section_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    section_p.paragraph_format.space_before = Pt(12)
    section_p.paragraph_format.space_after = Pt(6)
    run = section_p.add_run('各项决议内容：')
    set_run_font(run, size=14, bold=True)

    for idx, rd in enumerate(rd_data):
        code = rd.get('code', f'RD{idx+1:02d}')
        name = rd.get('name', '')
        resolution = rd.get('resolution', '经评审，该项目符合公司技术发展方向，技术方案可行，同意立项。')

        rd_title = doc.add_paragraph()
        rd_title.paragraph_format.space_before = Pt(10)
        run = rd_title.add_run(f'{idx+1}. {code} {name}')
        set_run_font(run, size=12, bold=True)

        add_normal_paragraph(doc, resolution, size=12)

    doc.add_paragraph()
    doc.add_paragraph()

    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sign_p.add_run(f'审批人(签字)：____________    日期：____________')
    set_run_font(run, size=12)

    company_p = doc.add_paragraph()
    company_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    company_p.paragraph_format.space_before = Pt(6)
    run = company_p.add_run(f'{enterprise}')
    set_run_font(run, size=12)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_p.add_run(f'{datetime.now().strftime("%Y年%m月%d日")}')
    set_run_font(run, size=12)

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
    doc.save(output_path)

    return {
        'output': output_path,
        'rd_count': len(rd_data),
        'codes': [rd.get('code', f'RD{i+1:02d}') for i, rd in enumerate(rd_data)],
    }


def main():
    parser = argparse.ArgumentParser(description='生成研发项目立项决议汇总文档')
    parser.add_argument('--rd-data', required=True, help='RD项目数据 JSON 文件路径')
    parser.add_argument('--enterprise', required=True, help='企业名称')
    parser.add_argument('--year', required=True, type=int, help='申报年度')
    parser.add_argument('--output', required=True, help='输出 .docx 文件路径')

    args = parser.parse_args()

    with open(args.rd_data, 'r', encoding='utf-8') as f:
        rd_data = json.load(f)

    if not isinstance(rd_data, list):
        print('Error: --rd-data must be a JSON array of RD project objects', file=sys.stderr)
        sys.exit(1)

    result = generate_approval_resolution(
        rd_data, args.enterprise, args.year, args.output
    )

    print(json.dumps(result, ensure_ascii=False))


if __name__ == '__main__':
    main()
