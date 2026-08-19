#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
template_filler.py - 模板填充引擎（XML级格式保持）

用法:
  python template_filler.py --template "path/to/template.docx" --data "path/to/data.json" --output "path/to/output.docx" [--verify] [--verbose]

核心目标: 填充内容时 100% 保持原模板格式。

技术要点:
  1. Document(template) 克隆 —— 不新建文档
  2. 遍历 paragraph.runs —— 匹配 {{placeholder}} 并替换，保持 run 格式不动
  3. 遍历 table.cells —— 同样匹配替换
  4. 合并单元格(vMerge/hMerge) —— XML层解除合并后填充
  5. safe_clip() —— 句号处安全截断
  6. --verify —— 后校验模板/产出样式差异
"""

import argparse
import json
import os
import re
import copy
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from lxml import etree


def safe_clip(text, max_chars):
    """安全截断：在完整句号处截断，放不下则整句跳过"""
    if len(text) <= max_chars:
        return text
    sentences = re.split(r'(?<=[。！？])', text)
    result = ''
    for s in sentences:
        if len(result) + len(s) <= max_chars:
            result += s
        else:
            break
    if not result:
        result = sentences[0][:max_chars]
    return result


def extract_placeholder(text):
    """从文本中提取 {{PLACEHOLDER}} 占位符名"""
    match = re.search(r'\{\{(\w+)\}\}', text)
    return match.group(1) if match else None


def find_placeholders_in_paragraph(paragraph):
    """在段落的所有 run 中查找占位符"""
    placeholders = []
    for run in paragraph.runs:
        if '{{' in run.text:
            matches = re.findall(r'\{\{(\w+)\}\}', run.text)
            for m in matches:
                placeholders.append((run, m))
    return placeholders


def find_placeholders_in_cell(cell):
    """在单元格的所有 paragraph 中查找占位符"""
    placeholders = []
    for paragraph in cell.paragraphs:
        placeholders.extend(find_placeholders_in_paragraph(paragraph))
    return placeholders


def replace_in_run(run, placeholder, new_text):
    """替换 run 中的占位符，保持 run 格式不变"""
    old_text = '{{' + placeholder + '}}'
    if old_text in run.text:
        run.text = run.text.replace(old_text, new_text)
        return True
    return False


def detect_vmerge(cell):
    """检测单元格是否有垂直合并"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        vmerge = tcPr.find(qn('w:vMerge'))
        if vmerge is not None:
            val = vmerge.get(qn('w:val'))
            return 'restart' if val == 'restart' else 'continue'
    return None


def detect_hmerge(cell):
    """检测单元格是否有水平合并"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        hmerge = tcPr.find(qn('w:hMerge'))
        if hmerge is not None:
            val = hmerge.get(qn('w:val'))
            return 'restart' if val == 'restart' else 'continue'
    return None


def unmerge_vertical(cell):
    """解除垂直合并"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        vmerge = tcPr.find(qn('w:vMerge'))
        if vmerge is not None:
            tcPr.remove(vmerge)


def analyze_template_structure(doc, verbose=False):
    """分析模板结构：打印所有表格/段落/嵌套层级"""
    info = {'paragraphs': 0, 'tables': 0, 'placeholders': [], 'merged_cells': []}

    info['paragraphs'] = len(doc.paragraphs)
    info['tables'] = len(doc.tables)

    for i, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                vmerge = detect_vmerge(cell)
                hmerge = detect_hmerge(cell)
                if vmerge or hmerge:
                    detail = f'Table[{i}] Row[{r}] Col[{c}]: vMerge={vmerge}, hMerge={hmerge}'
                    info['merged_cells'].append(detail)
                    if verbose:
                        print(f'  [MERGE] {detail}')

                placeholders_in_cell = find_placeholders_in_cell(cell)
                for run, placeholder in placeholders_in_cell:
                    detail = f'Table[{i}] Row[{r}] Col[{c}]: {{{{ {placeholder} }}}}'
                    info['placeholders'].append(detail)
                    if verbose:
                        print(f'  [PLACEHOLDER] {detail}')

                nested_tables = cell.tables
                if nested_tables and verbose:
                    for nt, ntbl in enumerate(nested_tables):
                        print(f'  [NESTED] Table[{i}] Row[{r}] Col[{c}] -> SubTable[{nt}] ({len(ntbl.rows)}r x {len(ntbl.columns)}c)')

    for p_idx, paragraph in enumerate(doc.paragraphs):
        placeholders_in_p = find_placeholders_in_paragraph(paragraph)
        for run, placeholder in placeholders_in_p:
            detail = f'Para[{p_idx}]: {{{{ {placeholder} }}}}'
            if detail not in info['placeholders']:
                info['placeholders'].append(detail)
            if verbose:
                print(f'  [PLACEHOLDER] {detail}')

    if verbose:
        print(f'\nStructure analysis: {info["paragraphs"]} paragraphs, {info["tables"]} tables, '
              f'{len(info["placeholders"])} placeholders, {len(info["merged_cells"])} merged cells')

    return info


def collect_style_info(doc):
    """收集文档的样式信息（用于 verify 对比）"""
    styles = []
    for p_idx, para in enumerate(doc.paragraphs):
        if para.runs:
            for r_idx, run in enumerate(para.runs):
                if run.text.strip():
                    font_info = {
                        'location': f'Para[{p_idx}].Run[{r_idx}]',
                        'text_preview': run.text[:50],
                        'font_name': run.font.name,
                        'font_size': str(run.font.size) if run.font.size else None,
                        'bold': run.font.bold,
                        'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                    }
                    styles.append(font_info)

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(para.runs):
                        if run.text.strip():
                            font_info = {
                                'location': f'Table[{t_idx}] Cell[{r_idx},{c_idx}] Para[{p_idx}].Run[{run_idx}]',
                                'text_preview': run.text[:50],
                                'font_name': run.font.name,
                                'font_size': str(run.font.size) if run.font.size else None,
                                'bold': run.font.bold,
                                'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                            }
                            styles.append(font_info)

    return styles


def verify_styles(template_styles, output_styles):
    """对比模板和产出样式差异"""
    issues = []
    template_map = {s['location']: s for s in template_styles}
    output_map = {s['location']: s for s in output_styles}

    common_locations = set(template_map.keys()) & set(output_map.keys())

    for loc in common_locations:
        t = template_map[loc]
        o = output_map[loc]
        if t['font_name'] != o['font_name']:
            issues.append(f'{loc}: font {t["font_name"]} -> {o["font_name"]}')
        if t['font_size'] != o['font_size']:
            issues.append(f'{loc}: size {t["font_size"]} -> {o["font_size"]}')
        if t['bold'] != o['bold']:
            issues.append(f'{loc}: bold {t["bold"]} -> {o["bold"]}')
        if t['color'] != o['color']:
            issues.append(f'{loc}: color {t["color"]} -> {o["color"]}')

    return issues


def fill_template(template_path, data, output_path, verbose=False, verify=False):
    """
    填充模板 — 核心函数

    template_path: 模板 .docx 路径
    data: dict, 占位符名 -> 替换文本
    output_path: 输出 .docx 路径
    verbose: 打印详细日志
    verify: 填充后进行样式对比校验
    """
    doc = Document(template_path)
    replaced_count = 0

    template_styles = collect_style_info(doc) if verify else None

    if verbose:
        print(f'Analyzing template: {template_path}')
        analyze_template_structure(doc, verbose=True)

    for paragraph in doc.paragraphs:
        placeholders = find_placeholders_in_paragraph(paragraph)
        for run, placeholder in placeholders:
            if placeholder in data:
                new_text = data[placeholder]
                if replace_in_run(run, placeholder, new_text):
                    replaced_count += 1
                    if verbose:
                        preview = new_text[:40] + '...' if len(new_text) > 40 else new_text
                        print(f'  Replaced {{{{ {placeholder} }}}} -> "{preview}"')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                vmerge = detect_vmerge(cell)
                if vmerge == 'restart':
                    unmerge_vertical(cell)
                    if verbose:
                        print(f'  Unmerged vertical merge cell')

                placeholders = find_placeholders_in_cell(cell)
                for run, placeholder in placeholders:
                    if placeholder in data:
                        new_text = data[placeholder]
                        if replace_in_run(run, placeholder, new_text):
                            replaced_count += 1
                            if verbose:
                                preview = new_text[:40] + '...' if len(new_text) > 40 else new_text
                                print(f'  Replaced {{{{ {placeholder} }}}} -> "{preview}"')

                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            nested_placeholders = find_placeholders_in_cell(nested_cell)
                            for run, placeholder in nested_placeholders:
                                if placeholder in data:
                                    new_text = data[placeholder]
                                    if replace_in_run(run, placeholder, new_text):
                                        replaced_count += 1
                                        if verbose:
                                            print(f'  [NESTED] Replaced {{{{ {placeholder} }}}}')

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    doc.save(output_path)

    result = {
        'output': output_path,
        'replaced_count': replaced_count,
        'total_placeholders': len([k for k in data if k]),
    }

    if verify:
        output_styles = collect_style_info(doc)
        style_issues = verify_styles(template_styles, output_styles)
        result['style_issues'] = style_issues
        result['style_consistency'] = 1.0 - len(style_issues) / max(len(template_styles), 1)

    return result


def fill_template_batch(template_path, data_list, output_dir, prefix='RD', verbose=False, verify=False):
    """
    批量填充模板 — 一次模板多次填充

    template_path: 模板 .docx 路径
    data_list: list[dict], 每个 dict 对应一份 RD 报告的数据
    output_dir: 输出目录
    prefix: 输出文件名前缀
    """
    results = []
    for i, data in enumerate(data_list):
        code = data.get('RD_CODE', f'{prefix}{i+1:02d}')
        output_path = os.path.join(output_dir, f'{code}_{data.get("RD_NAME", "project")[:20]}.docx')
        result = fill_template(template_path, data, output_path, verbose=verbose, verify=verify)
        results.append(result)
        if verbose:
            print(f'  [{i+1}/{len(data_list)}] Generated: {output_path}')
    return results


def main():
    parser = argparse.ArgumentParser(description='RD立项报告模板填充引擎')
    parser.add_argument('--template', required=True, help='模板 .docx 文件路径')
    parser.add_argument('--data', required=True, help='填充数据 JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出 .docx 文件路径')
    parser.add_argument('--verify', action='store_true', help='填充后进行样式对比校验')
    parser.add_argument('--verbose', '-v', action='store_true', help='打印详细日志')
    parser.add_argument('--batch', action='store_true', help='批量模式：data 为 list 时批量生成')
    parser.add_argument('--output-dir', help='批量模式的输出目录')

    args = parser.parse_args()

    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)

    if args.batch:
        if not args.output_dir:
            print('Error: --batch requires --output-dir', file=sys.stderr)
            sys.exit(1)
        if not isinstance(data, list):
            print('Error: --batch requires data to be a JSON array', file=sys.stderr)
            sys.exit(1)

        results = fill_template_batch(
            args.template, data, args.output_dir,
            verbose=args.verbose, verify=args.verify
        )

        total_replaced = sum(r['replaced_count'] for r in results)

        if args.verify:
            total_issues = sum(len(r.get('style_issues', [])) for r in results)
            avg_consistency = sum(r.get('style_consistency', 0) for r in results) / max(len(results), 1)
            print(f'Batch complete: {len(results)} files, {total_replaced} replacements, '
                  f'{total_issues} style issues, avg consistency: {avg_consistency:.1%}')

        print(json.dumps({
            'mode': 'batch',
            'files': len(results),
            'total_replaced': total_replaced,
        }, ensure_ascii=False))
    else:
        if not isinstance(data, dict):
            print('Error: --data must be a JSON object for single mode', file=sys.stderr)
            sys.exit(1)

        result = fill_template(
            args.template, data, args.output,
            verbose=args.verbose, verify=args.verify
        )

        if args.verify:
            print(f'Style consistency: {result["style_consistency"]:.1%} '
                  f'({len(result.get("style_issues", []))} issues)')

        print(json.dumps({
            'output': result['output'],
            'replaced_count': result['replaced_count'],
        }, ensure_ascii=False))


if __name__ == '__main__':
    main()
