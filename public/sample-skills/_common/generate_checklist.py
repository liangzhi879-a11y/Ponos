"""资料收集清单生成工具（独立脚本 v1.16.0）

解决"agent只输出xlsx不生成docx清单"问题。
agent 准备好清单数据 JSON 后，通过此脚本生成标准格式的 docx 收资清单。

v1.16.0 新增：
  - 章节：2025年研发项目及成果转化补缺清单（基于加计扣除表对比）
  - 章节：按时间节点需收集合同清单（无合同编号则列发票编号）
  - 章节：2025年优质合同发票推荐收集清单（基于全量发票筛选）
  - 合同发票清单冗余系数标注（10-20%多收集）

用法：
  python generate_checklist.py --enterprise "深圳市XX公司" --year 2026 \\
    --data checklist_data.json --output-dir "07_资料收集清单"

数据JSON格式（checklist_data.json）：
{
  "enterprise_name": "深圳市XX公司",
  "application_year": 2026,
  "recent_three_years": [2023, 2024, 2025],
  "last_year": 2025,
  "items": [
    {
      "department": "行政", "category": "基础材料", "seq": 1,
      "name": "企业最近的营业执照", "provide_way": "扫描件",
      "requirement": "需盖贵司公章", "deadline": "2026-08-15",
      "provided": "否", "progress": "", "remark": ""
    },
    ...
  ],
  "equipment_list": [
    {"seq": 1, "name": "设备名", "asset_code": "XX-001", "dept": "研发部",
     "purchase_date": "2024-03-01", "value": 5.8, "category": "研发设备", "confirmed": "✗"}
  ],
  "contract_invoice_by_year": {
    "2023": [...], "2024": [...], "2025": [...]
  },
  "supplement_summary": {
    "total_items": 50, "provided_count": 20, "pending_count": 30,
    "by_dept": {"行政": 5, "人事": 10, "财务": 8, "研发": 7}
  },
  "missing_rd_2025": {
    "source_file": "2026加计扣除表.xlsx",
    "items": [
      {"rd_code": "RD21", "rd_name": "2025年新立项研发项目名称",
       "start_date": "2025-03-01", "end_date": "2025-12-31",
       "rd_budget": 50.0, "linked_ip": "", "linked_ach": "",
       "source": "2026加计扣除表行15", "remark": "核心表格需补充"}
    ]
  },
  "contracts_needed_by_timepoint": {
    "timepoints": ["2024-03-15", "2024-06-01", "2025-01-20"],
    "items": [
      {"seq": 1, "timepoint": "2024-03-15", "linked_rd": "RD01",
       "customer": "深圳XX公司", "amount": 10.5,
       "has_contract": false, "has_invoice": true,
       "invoice_no": "FP20240315001", "contract_no": "",
       "status": "缺合同，需补充", "remark": "无合同编号，用发票编号代替"}
    ]
  },
  "recommended_invoices_2025": {
    "source_file": "2025全量发票.xlsx",
    "filter_criteria": {"min_amount": 5.0, "date_range": "2025-01~2025-12"},
    "redundancy_ratio": 0.15,
    "items": [
      {"seq": 1, "invoice_no": "FP20250315001", "invoice_date": "2025-03-15",
       "customer": "深圳XX公司", "customer_quality": "优质",
       "amount": 25.8, "has_contract": false, "reason": "金额大+客户优质+对应RD时间"}
    ]
  }
}

依赖：python-docx
"""
import os
import sys
import json
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell):
    """设置单元格边框（thin）"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_run_font(run, name='宋体', size=11, bold=False):
    """设置run字体（中文宋体+英文Times New Roman）"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


def add_heading_with_style(doc, text, level=2):
    """添加Heading样式的段落（生成文档大纲导航）"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '宋体'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
    return heading


def add_table_with_header(doc, headers, rows, col_widths=None):
    """添加带表头的表格（应用thin边框+表头蓝底白字）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头行
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(header))
        set_run_font(run, size=10, bold=True)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # 白字
        set_cell_shading(cell, '4472C4')  # 蓝底
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(value) if value is not None else '')
            set_run_font(run, size=10)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def generate_checklist_docx(data, output_path):
    """生成收资清单docx

    文档结构：
      1. 标题 + 企业信息
      2. 整理概览（已收到X份、仍缺Y份、按部门统计）
      3. 主清单（10列三层分类：部门/类别/序号/资料名/提供方式/要求/截止日期/是否已提供/进度/备注）
      4. 研发设备清单（8列）
      5. 合同发票清单（9列，按年度分3个表）
      6. 补充需求说明
    """
    doc = Document()

    # 设置全局字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')

    enterprise = data.get('enterprise_name', '')
    year = data.get('application_year', 2026)
    recent = data.get('recent_three_years', [year-3, year-2, year-1])
    last_year = data.get('last_year', year - 1)

    # 1. 标题
    title = doc.add_heading(f'{enterprise}-{year}年高新认定资料收集清单', level=0)
    for run in title.runs:
        set_run_font(run, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}    '
                    f'近三年：{recent[0]}-{recent[-1]}    上年度：{last_year}')
    set_run_font(run, size=10)

    doc.add_paragraph('')

    # 2. 整理概览
    add_heading_with_style(doc, '一、整理概览', level=1)
    summary = data.get('supplement_summary', {})
    total = summary.get('total_items', len(data.get('items', [])))
    provided = summary.get('provided_count', 0)
    pending = summary.get('pending_count', total - provided)
    by_dept = summary.get('by_dept', {})

    p = doc.add_paragraph()
    run = p.add_run(f'共 {total} 项资料，已收到 {provided} 项，仍需补充 {pending} 项。')
    set_run_font(run, size=11, bold=True)

    if by_dept:
        p = doc.add_paragraph()
        run = p.add_run('按部门统计：')
        set_run_font(run, size=11, bold=True)
        for dept, cnt in by_dept.items():
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f'{dept}：{cnt} 项')
            set_run_font(run, size=11)

    doc.add_paragraph('')

    # 3. 主清单（10列）
    add_heading_with_style(doc, '二、资料收集清单（10列三层分类）', level=1)

    headers = ['负责部门', '类别', '序号', '所需资料', '提供方式',
               '材料要求说明', '计划完成时间', '是否已经提供',
               '资料整理及撰写进度情况', '备注']
    col_widths = [1.5, 1.5, 0.8, 3.0, 1.5, 3.5, 1.5, 1.2, 2.5, 1.5]

    items = data.get('items', [])
    rows = []
    for item in items:
        rows.append([
            item.get('department', ''),
            item.get('category', ''),
            item.get('seq', ''),
            item.get('name', ''),
            item.get('provide_way', ''),
            item.get('requirement', ''),
            item.get('deadline', ''),
            item.get('provided', '否'),
            item.get('progress', ''),
            item.get('remark', ''),
        ])

    add_table_with_header(doc, headers, rows, col_widths)
    doc.add_paragraph('')

    # 4. 研发设备清单（8列）
    equipment = data.get('equipment_list', [])
    if equipment:
        add_heading_with_style(doc, '三、研发设备清单（8列）', level=1)
        equip_headers = ['序号', '设备名称', '资产编码', '部门',
                         '购置日期', '原值（万元）', '分类', '核对确认']
        equip_widths = [1.0, 3.0, 2.0, 1.5, 1.5, 1.5, 1.5, 1.5]
        equip_rows = []
        for eq in equipment:
            equip_rows.append([
                eq.get('seq', ''),
                eq.get('name', ''),
                eq.get('asset_code', ''),
                eq.get('dept', ''),
                eq.get('purchase_date', ''),
                eq.get('value', ''),
                eq.get('category', '研发设备'),
                eq.get('confirmed', '✗'),
            ])
        add_table_with_header(doc, equip_headers, equip_rows, equip_widths)
        doc.add_paragraph('')

    # 5. 合同发票清单（9列，按年度分3个表）
    contract_invoice = data.get('contract_invoice_by_year', {})
    if contract_invoice:
        add_heading_with_style(doc, '四、合同发票清单（9列，按年度分表）', level=1)
        # 冗余系数提示
        redundancy_ratio = data.get('redundancy_ratio', 0.15)
        if redundancy_ratio:
            p = doc.add_paragraph()
            run = p.add_run(f'⚠ 冗余收集要求：发票合同数量需多收集 {int(redundancy_ratio*100)}%（即在原筛选基础上增加 {int(redundancy_ratio*100)}% 的冗余量，用于应对审核补料）')
            set_run_font(run, size=11, bold=True)

        ci_headers = ['序号', '合同编号', '合同名称', '客户名称',
                      '合同金额（万元）', '合同签订日期', '发票号码',
                      '发票金额（万元）', '状态']
        ci_widths = [0.8, 1.5, 3.0, 2.0, 1.5, 1.5, 1.8, 1.5, 1.2]

        for year_key in sorted(contract_invoice.keys()):
            ci_items = contract_invoice[year_key]
            if not ci_items:
                continue
            add_heading_with_style(doc, f'{year_key}年度合同发票清单', level=2)
            ci_rows = []
            for ci in ci_items:
                has_contract = '合同编号' in ci and ci['合同编号']
                has_invoice = '发票号码' in ci and ci['发票号码']
                if has_contract and has_invoice:
                    status = '合同已有✓ 发票已有✓'
                elif has_contract:
                    status = '合同已有✓ 缺发票，需补充'
                elif has_invoice:
                    status = '缺合同，需补充 发票已有✓'
                else:
                    status = '缺合同，需补充 缺发票，需补充'
                ci_rows.append([
                    ci.get('seq', ''),
                    ci.get('合同编号', ''),
                    ci.get('合同名称', ''),
                    ci.get('客户名称', ''),
                    ci.get('合同金额', ''),
                    ci.get('合同签订日期', ''),
                    ci.get('发票号码', ''),
                    ci.get('发票金额', ''),
                    status,
                ])
            add_table_with_header(doc, ci_headers, ci_rows, ci_widths)

            # 【补充需求】合并行
            missing_count = sum(1 for r in ci_rows if '缺' in r[-1])
            target_count = int(len(ci_rows) * (1 + redundancy_ratio)) if redundancy_ratio else len(ci_rows)
            if missing_count > 0 or redundancy_ratio:
                p = doc.add_paragraph()
                run = p.add_run(f'【补充需求】{year_key}年度共需补充 {missing_count} 项（合同或发票缺失项）'
                                f'  |  建议收集总量 {target_count} 项（含 {int(redundancy_ratio*100)}% 冗余）')
                set_run_font(run, size=11, bold=True)
            doc.add_paragraph('')

    # 6. 2025年研发项目及成果转化补缺清单
    missing_rd = data.get('missing_rd_2025', {})
    if missing_rd and missing_rd.get('items'):
        add_heading_with_style(doc, '五、2025年研发项目及成果转化补缺清单', level=1)
        source_file = missing_rd.get('source_file', '2026加计扣除表')
        p = doc.add_paragraph()
        run = p.add_run(f'数据来源：{source_file}（识别2025年新立项但未纳入核心表格的RD）')
        set_run_font(run, size=11, bold=True)

        rd_headers = ['序号', 'RD编号', '研发项目名称', '立项时间', '结束时间',
                      '研发预算（万元）', '关联IP', '关联成果转化', '数据来源', '备注']
        rd_widths = [0.8, 1.2, 3.5, 1.5, 1.5, 1.5, 1.2, 1.5, 1.8, 2.0]
        rd_rows = []
        for idx, item in enumerate(missing_rd['items'], 1):
            rd_rows.append([
                idx,
                item.get('rd_code', ''),
                item.get('rd_name', ''),
                item.get('start_date', ''),
                item.get('end_date', ''),
                item.get('rd_budget', ''),
                item.get('linked_ip', ''),
                item.get('linked_ach', ''),
                item.get('source', ''),
                item.get('remark', '核心表格需补充'),
            ])
        add_table_with_header(doc, rd_headers, rd_rows, rd_widths)

        p = doc.add_paragraph()
        run = p.add_run(f'⚠ 以上 {len(rd_rows)} 项2025年研发项目需补充到核心表格（RD表+成果转化表）')
        set_run_font(run, size=11, bold=True)
        doc.add_paragraph('')

    # 7. 按时间节点需收集合同清单
    contracts_needed = data.get('contracts_needed_by_timepoint', {})
    if contracts_needed and contracts_needed.get('items'):
        add_heading_with_style(doc, '六、按时间节点需收集合同清单（无合同编号则列发票编号）', level=1)
        timepoints = contracts_needed.get('timepoints', [])
        if timepoints:
            p = doc.add_paragraph()
            run = p.add_run(f'时间节点：{", ".join(timepoints)}')
            set_run_font(run, size=11, bold=True)

        cn_headers = ['序号', '时间节点', '关联RD/PS', '客户名称', '金额（万元）',
                      '合同编号', '发票号码', '状态', '备注']
        cn_widths = [0.8, 1.5, 1.2, 2.5, 1.5, 1.8, 1.8, 1.8, 2.0]
        cn_rows = []
        for item in contracts_needed['items']:
            cn_rows.append([
                item.get('seq', ''),
                item.get('timepoint', ''),
                item.get('linked_rd', '') or item.get('linked_ps', ''),
                item.get('customer', ''),
                item.get('amount', ''),
                item.get('contract_no', ''),
                item.get('invoice_no', ''),
                item.get('status', ''),
                item.get('remark', ''),
            ])
        add_table_with_header(doc, cn_headers, cn_rows, cn_widths)

        missing_contract_count = sum(1 for r in cn_rows if '缺合同' in r[-2])
        p = doc.add_paragraph()
        run = p.add_run(f'⚠ 共需补充 {missing_contract_count} 份合同（已列发票编号可代替合同编号用于核对）')
        set_run_font(run, size=11, bold=True)
        doc.add_paragraph('')

    # 8. 2025年优质合同发票推荐收集清单
    recommended = data.get('recommended_invoices_2025', {})
    if recommended and recommended.get('items'):
        add_heading_with_style(doc, '七、2025年优质合同发票推荐收集清单', level=1)
        source_file = recommended.get('source_file', '2025全量发票.xlsx')
        criteria = recommended.get('filter_criteria', {})
        redundancy_ratio = recommended.get('redundancy_ratio', 0.15)
        p = doc.add_paragraph()
        run = p.add_run(f'数据来源：{source_file}')
        set_run_font(run, size=11, bold=True)

        if criteria:
            p = doc.add_paragraph()
            run = p.add_run(f'筛选条件：{criteria}')
            set_run_font(run, size=10)

        p = doc.add_paragraph()
        run = p.add_run(f'冗余系数：{int(redundancy_ratio*100)}%（在筛选结果基础上多收集 {int(redundancy_ratio*100)}%）')
        set_run_font(run, size=11, bold=True)

        ri_headers = ['序号', '发票号码', '开票日期', '客户名称', '客户优质度',
                      '金额（万元）', '是否有合同', '推荐理由']
        ri_widths = [0.8, 2.0, 1.5, 2.5, 1.5, 1.5, 1.2, 3.5]
        ri_rows = []
        for item in recommended['items']:
            ri_rows.append([
                item.get('seq', ''),
                item.get('invoice_no', ''),
                item.get('invoice_date', ''),
                item.get('customer', ''),
                item.get('customer_quality', ''),
                item.get('amount', ''),
                '是✓' if item.get('has_contract') else '否✗',
                item.get('reason', ''),
            ])
        add_table_with_header(doc, ri_headers, ri_rows, ri_widths)

        no_contract_count = sum(1 for r in ri_rows if '否' in r[-2])
        target_count = int(len(ri_rows) * (1 + redundancy_ratio))
        p = doc.add_paragraph()
        run = p.add_run(f'⚠ 推荐收集 {len(ri_rows)} 项（含 {no_contract_count} 项需补合同），'
                        f'考虑 {int(redundancy_ratio*100)}% 冗余后建议收集总量 {target_count} 项')
        set_run_font(run, size=11, bold=True)
        doc.add_paragraph('')

    # 9. 补充需求说明
    add_heading_with_style(doc, '八、补充需求说明', level=1)
    if pending > 0:
        p = doc.add_paragraph()
        run = p.add_run(f'本次共需补充 {pending} 项资料，请各部门按照上述清单在计划完成时间内提供。')
        set_run_font(run, size=11)
    else:
        p = doc.add_paragraph()
        run = p.add_run('所有资料已齐全，无需补充。')
        set_run_font(run, size=11)

    # 保存
    doc.save(output_path)
    print(f'[生成] 收资清单docx: {output_path}')
    print(f'  主清单: {len(items)} 项')
    print(f'  设备清单: {len(equipment)} 项')
    print(f'  合同发票: {sum(len(v) for v in contract_invoice.values())} 项')
    if missing_rd and missing_rd.get('items'):
        print(f'  2025年RD补缺: {len(missing_rd["items"])} 项')
    if contracts_needed and contracts_needed.get('items'):
        print(f'  按时间节点需补合同: {len(contracts_needed["items"])} 项')
    if recommended and recommended.get('items'):
        print(f'  2025年推荐发票: {len(recommended["items"])} 项')
    return output_path


def validate_checklist_quality(docx_path):
    """校验docx清单质量（10项指标 v1.16.0）"""
    from docx import Document
    doc = Document(docx_path)
    checks = {
        'main_table_10_cols': False,   # 主清单10列表格存在
        'equipment_table_8_cols': False, # 设备表8列
        'contract_invoice_9_cols': False, # 合同发票9列
        'status_with_checkmark': False, # 状态标识含✓
        'heading_styles': False,        # Heading样式存在
        'supplement_summary': False,    # 量化描述存在
        'supplement_demand_row': False, # 【补充需求】行存在
        'missing_rd_section': False,    # 2025年RD补缺章节存在
        'contracts_needed_section': False, # 按时间节点需收集合同章节存在
        'recommended_invoices_section': False, # 推荐发票章节存在
    }

    full_text = '\n'.join(p.text for p in doc.paragraphs)

    for table in doc.tables:
        cols = len(table.columns)
        if cols == 10:
            checks['main_table_10_cols'] = True
        elif cols == 8:
            checks['equipment_table_8_cols'] = True
        elif cols == 9:
            checks['contract_invoice_9_cols'] = True
        elif cols == 10:  # RD补缺表也是10列，但通过章节标题区分
            pass

        # 检查状态标识
        for row in table.rows:
            for cell in row.cells:
                if '✓' in cell.text or '✗' in cell.text:
                    checks['status_with_checkmark'] = True
                    break

    # 检查Heading样式
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            checks['heading_styles'] = True
            break

    # 检查量化描述和补充需求
    if '已收到' in full_text and '项' in full_text:
        checks['supplement_summary'] = True
    if '【补充需求】' in full_text:
        checks['supplement_demand_row'] = True

    # 检查3个新章节
    if '2025年研发项目及成果转化补缺清单' in full_text:
        checks['missing_rd_section'] = True
    if '按时间节点需收集合同清单' in full_text:
        checks['contracts_needed_section'] = True
    if '优质合同发票推荐收集清单' in full_text:
        checks['recommended_invoices_section'] = True

    all_passed = all(checks.values())
    return {
        'all_passed': all_passed,
        'checks': checks,
        'report': '\n'.join(f'  {"✓" if v else "✗"} {k}' for k, v in checks.items()),
    }


def main():
    parser = argparse.ArgumentParser(description='资料收集清单生成工具 v1.16.0')
    parser.add_argument('--enterprise', required=True, help='企业名称')
    parser.add_argument('--year', type=int, default=2026, help='申报年份')
    parser.add_argument('--data', required=True, help='清单数据JSON文件路径')
    parser.add_argument('--output-dir', required=True, help='输出目录')
    parser.add_argument('--output-name', default=None, help='输出文件名（不含路径）')
    args = parser.parse_args()

    # 加载数据
    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 补充字段
    data.setdefault('enterprise_name', args.enterprise)
    data.setdefault('application_year', args.year)
    data.setdefault('recent_three_years', [args.year-3, args.year-2, args.year-1])
    data.setdefault('last_year', args.year - 1)

    # 输出路径
    os.makedirs(args.output_dir, exist_ok=True)
    output_name = args.output_name or f'{args.enterprise}-{args.year}年高新认定资料收集清单.docx'
    output_path = os.path.join(args.output_dir, output_name)

    # 生成
    generate_checklist_docx(data, output_path)

    # 校验
    print('\n[校验] 质量校验（10项指标 v1.16.0）:')
    result = validate_checklist_quality(output_path)
    print(result['report'])
    if result['all_passed']:
        print('\n[校验] ✓ 全部通过')
    else:
        print('\n[校验] ✗ 存在不通过项，需检查')
        # v1.16.0：3个新章节未提供数据时不视为失败（可选章节）
        optional_checks = ['missing_rd_section', 'contracts_needed_section', 'recommended_invoices_section']
        failed_required = [k for k, v in result['checks'].items() if not v and k not in optional_checks]
        if not failed_required:
            print('[校验] ⚠ 仅可选章节（RD补缺/合同清单/推荐发票）未提供数据，其他必选项全部通过')
        else:
            sys.exit(1)


if __name__ == '__main__':
    main()
