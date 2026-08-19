#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
file_content_classifier.py - 基于文件内容的资料分类器

通过读取文件内容（非仅文件名）自动判定高新技术企业认定项目中的资料类型。

用法：
  python file_content_classifier.py scan --dir "资料目录" --output "分类结果.json" --exclude-dir ".trae"
  python file_content_classifier.py filter --input "分类结果.json" --year 2026 --target-dir "整理输出目录" --project-root "项目根目录" --copy
"""

import argparse
import json
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

sys.path.insert(0, str(Path(__file__).parent))


CLASSIFICATION_PATTERNS = {
    "专利证书": {
        "category_id": "1",
        "keywords": ["发明专利证书", "实用新型专利证书", "外观设计专利证书", "证书号",
                      "专利号", "专利权人", "授权公告日"],
        "date_patterns": [r"授权公告日[：:]\s*(\d{4})", r"申请日[：:]\s*(\d{4})"],
        "weight": 1.2,
    },
    "软著证书": {
        "category_id": "1",
        "keywords": ["计算机软件著作权登记证书", "软件著作权", "登记号", "软件名称"],
        "date_patterns": [r"开发完成日期[：:]\s*(\d{4})"],
        "weight": 1.2,
    },
    "立项报告": {
        "category_id": "2",
        "keywords": ["研发立项报告", "研究开发项目", "项目编号", "项目负责人",
                      "起止时间", "研发项目", "立项申请", "项目验收", "结题报告", "任务书"],
        "date_patterns": [r"(\d{4})年[度]*[^\n]{0,20}研发", r"起止时间[：:]\s*(\d{4})"],
        "weight": 1.1,
    },
    "审计报告": {
        "category_id": "8",
        "keywords": ["审计报告", "注册会计师", "审计意见", "财务报表",
                      "资产负债表", "利润表", "现金流量表"],
        "date_patterns": [r"(\d{4})年度审计", r"审计年度[：:]\s*(\d{4})"],
        "weight": 1.1,
    },
    "纳税申报表": {
        "category_id": "9",
        "keywords": ["企业所得税", "纳税申报", "应纳税所得额",
                      "企业所得税年度纳税申报表", "中华人民共和国企业所得税"],
        "date_patterns": [r"(\d{4})年度.*所得税", r"税款所属期间.*?(\d{4})"],
        "weight": 1.1,
    },
    "合同": {
        "category_id": "17",
        "keywords": ["合同", "协议", "甲方", "乙方", "签订日期", "合同编号"],
        "date_patterns": [r"签订日期[：:]\s*(\d{4})", r"签订时间[：:]\s*(\d{4})"],
        "weight": 1.0,
    },
    "发票": {
        "category_id": "17",
        "keywords": ["发票", "发票代码", "发票号码", "开票日期",
                      "销售方", "购买方", "货物或应税劳务"],
        "date_patterns": [r"开票日期[：:]\s*(\d{4})"],
        "weight": 1.0,
    },
    "营业执照": {
        "category_id": "6",
        "keywords": ["营业执照", "统一社会信用代码", "法定代表人",
                      "注册资本", "经营范围", "成立日期"],
        "date_patterns": [],
        "weight": 1.3,
    },
    "社保": {
        "category_id": "16",
        "keywords": ["社会保险", "社保", "缴费", "参保",
                      "养老保险", "医疗保险", "工伤保险"],
        "date_patterns": [r"(\d{4})年.*社保", r"(\d{4})年\d{1,2}月"],
        "weight": 1.0,
    },
    "检测报告": {
        "category_id": "19",
        "keywords": ["检测报告", "检验报告", "CMA", "CNAS",
                      "检测结果", "检验检测"],
        "date_patterns": [r"检测日期[：:]\s*(\d{4})", r"报告日期[：:]\s*(\d{4})"],
        "weight": 1.0,
    },
    "管理制度": {
        "category_id": "12",
        "keywords": ["管理制度", "管理办法", "研发制度", "组织管理",
                      "产学研", "激励制度", "辅助账", "核算办法"],
        "date_patterns": [],
        "weight": 0.9,
    },
    "学历证书": {
        "category_id": "16",
        "keywords": ["毕业证书", "学位证书", "学历", "学位", "毕业于"],
        "date_patterns": [],
        "weight": 1.0,
    },
    "设备清单": {
        "category_id": "19",
        "keywords": ["设备清单", "仪器设备", "固定资产",
                      "设备名称", "规格型号", "购置日期"],
        "date_patterns": [],
        "weight": 1.0,
    },
    "花名册": {
        "category_id": "16",
        "keywords": ["花名册", "员工名册", "职工名册", "人员名单",
                      "员工信息", "部门", "职务", "入职日期"],
        "date_patterns": [],
        "weight": 1.0,
    },
    "查新报告": {
        "category_id": "19",
        "keywords": ["科技查新", "查新报告", "查新项目", "查新点"],
        "date_patterns": [r"查新完成日期[：:]\s*(\d{4})"],
        "weight": 1.1,
    },
    "完税证明": {
        "category_id": "9",
        "keywords": ["完税证明", "税收完税证明", "纳税凭证", "完税凭证"],
        "date_patterns": [r"(\d{4})年.*完税"],
        "weight": 1.0,
    },
    "产品手册": {
        "category_id": "3",
        "keywords": ["产品手册", "产品介绍", "产品规格", "产品说明",
                      "技术参数", "产品型号"],
        "date_patterns": [],
        "weight": 0.9,
    },
    "照片": {
        "category_id": "19",
        "keywords": [],
        "is_image_only": True,
        "date_patterns": [],
        "weight": 0.5,
    },
}


FILE_ORGANIZER_CATEGORIES = {
    '1':  {'name': '1.IP证明材料（知识产权）'},
    '2':  {'name': '2.企业研究开发活动情况证明材料-立项报告任务书验收报告（RD）'},
    '3':  {'name': '3.PS证明材料（高新技术产品）'},
    '4':  {'name': '4.科技成果转化'},
    '5':  {'name': '5.标准资料（参与制定的各类标准文件）'},
    '6':  {'name': '6.营业执照'},
    '8':  {'name': '8.前三年财务审计报告'},
    '9':  {'name': '9.前三年企业所得税纳税申报表'},
    '10': {'name': '10.前三年研发费用专项审计报告'},
    '11': {'name': '11.上年度高新产品（服务）收入专项审计报告'},
    '12': {'name': '12-15.组织管理制度'},
    '16': {'name': '16.人力资源情况证明材料'},
    '17': {'name': '17.上年度与高新技术产品（服务）相关的代表性的销售合同与发票'},
    '18': {'name': '18.往期项目资料'},
    '19': {'name': '19.补充资料'},
    '98': {'name': '98.历史参考资料'},
    '99': {'name': '99.其他资料'},
}


SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.xls', '.jpg', '.jpeg', '.png', '.txt'}


def is_image_file(filepath):
    return filepath.suffix.lower() in {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif', '.webp'}


# ============================================================
# 内容提取函数
# ============================================================

def extract_pdf_text(filepath, max_pages=5):
    import fitz
    texts = []
    try:
        doc = fitz.open(str(filepath))
        total_pages = len(doc)
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            page_text = page.get_text()
            if page_text.strip():
                texts.append(page_text)
        doc.close()
        combined = "\n".join(texts)
        if len(combined.strip()) < 50:
            return None
        return combined
    except Exception:
        return None


def extract_pdf_text_pdfplumber(filepath, max_pages=5):
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(str(filepath)) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                if i >= max_pages:
                    break
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
        combined = "\n".join(texts)
        if len(combined.strip()) < 50:
            return None
        return combined
    except Exception:
        return None


def extract_pdf_text_pypdf2(filepath, max_pages=5):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(filepath))
        texts = []
        for i, page in enumerate(reader.pages):
            if i >= max_pages:
                break
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text)
        combined = "\n".join(texts)
        if len(combined.strip()) < 50:
            return None
        return combined
    except Exception:
        return None


def extract_pdf_via_ocr(filepath, project_name="default"):
    try:
        from ocr_engine import OCREngine
        engine = OCREngine()
        result = engine.ocr_pdf(str(filepath), force_ocr=False, use_cache=True, project_name=project_name)
        text = result.get("text", "")
        if len(text.strip()) < 10:
            return None
        return text
    except Exception:
        return None


def read_pdf_content(filepath, project_name="default"):
    text = extract_pdf_text(filepath, max_pages=5)
    if text:
        return text
    text = extract_pdf_text_pdfplumber(filepath, max_pages=5)
    if text:
        return text
    text = extract_pdf_text_pypdf2(filepath, max_pages=5)
    if text:
        return text
    text = extract_pdf_via_ocr(filepath, project_name=project_name)
    if text:
        return text
    return ""


def read_docx_content(filepath):
    try:
        import docx
        doc = docx.Document(str(filepath))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception:
        return ""


def read_xlsx_content(filepath, max_rows=5):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(filepath), data_only=True, read_only=True)
        parts = []
        parts.append("工作表: " + ", ".join(wb.sheetnames))
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"\n--- {sheet_name} ---")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    break
                row_vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in row_vals):
                    parts.append(" | ".join(row_vals))
        wb.close()
        return "\n".join(parts)
    except Exception:
        return ""


def read_xls_content(filepath, max_rows=5):
    try:
        import xlrd
        wb = xlrd.open_workbook(str(filepath))
        parts = []
        parts.append("工作表: " + ", ".join(wb.sheet_names()))
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            parts.append(f"\n--- {sheet_name} ---")
            for i in range(min(ws.nrows, max_rows)):
                row_vals = [str(ws.cell_value(i, j)).strip() for j in range(ws.ncols)]
                if any(v for v in row_vals):
                    parts.append(" | ".join(row_vals))
        return "\n".join(parts)
    except Exception:
        return ""


def read_image_via_ocr(filepath, project_name="default"):
    try:
        from ocr_engine import OCREngine
        engine = OCREngine()
        result = engine.ocr_image(str(filepath), project_name=project_name)
        text = result.get("text", "")
        if len(text.strip()) < 5:
            return None
        return text
    except Exception:
        return None


def read_txt_content(filepath):
    try:
        return filepath.read_text(encoding='utf-8')
    except UnicodeDecodeError:
        try:
            return filepath.read_text(encoding='gbk')
        except Exception:
            return ""


def extract_file_content(filepath, project_name="default"):
    ext = filepath.suffix.lower()
    if ext == '.pdf':
        return read_pdf_content(filepath, project_name=project_name)
    elif ext == '.docx':
        return read_docx_content(filepath)
    elif ext == '.xlsx':
        return read_xlsx_content(filepath)
    elif ext == '.xls':
        return read_xls_content(filepath)
    elif is_image_file(filepath):
        return read_image_via_ocr(filepath, project_name=project_name) or ""
    elif ext == '.txt':
        return read_txt_content(filepath)
    return ""


# ============================================================
# 分类引擎
# ============================================================

def extract_dates_from_text(text, patterns):
    dates = []
    for pattern in patterns:
        for m in re.finditer(pattern, text):
            year_str = m.group(1)
            try:
                year = int(year_str)
                if 1980 <= year <= 2040:
                    dates.append(year)
            except ValueError:
                pass
    return sorted(set(dates))


def extract_all_years_from_text(text):
    years = set()
    for m in re.finditer(r'(?<!\d)(20\d{2})(?!\d)', text):
        try:
            year = int(m.group(1))
            if 1980 <= year <= 2040:
                years.add(year)
        except ValueError:
            pass
    return sorted(years)


def classify_single_file(filepath, project_name="default"):
    ext = filepath.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS and not is_image_file(filepath):
        return {
            "file": str(filepath),
            "file_type": ext,
            "classification": "其他",
            "category_id": "99",
            "confidence": 1.0,
            "matched_keywords": [],
            "extracted_dates": [],
            "needs_review": True,
            "content_preview": "",
            "error": f"不支持的文件类型: {ext}",
        }

    content = extract_file_content(filepath, project_name=project_name)
    content_preview = (content or "")[:500]

    if is_image_file(filepath) and not content:
        return {
            "file": str(filepath),
            "file_type": ext.lstrip('.'),
            "classification": "照片",
            "category_id": "19",
            "confidence": 0.5,
            "matched_keywords": [],
            "extracted_dates": [],
            "needs_review": False,
            "content_preview": content_preview,
            "error": "",
        }

    best_type = "其他"
    best_category_id = "99"
    best_score = 0.0
    best_matched = []
    best_dates = []
    best_weight = 1.0

    for type_name, cfg in CLASSIFICATION_PATTERNS.items():
        if cfg.get("is_image_only"):
            continue

        keywords = cfg.get("keywords", [])
        weight = cfg.get("weight", 1.0)
        date_patterns = cfg.get("date_patterns", [])

        if not keywords:
            continue

        matched = []
        for kw in keywords:
            if kw in content:
                matched.append(kw)

        kw_ratio = len(matched) / len(keywords) if keywords else 0
        score = kw_ratio * weight

        if score > best_score:
            best_score = score
            best_type = type_name
            best_category_id = cfg.get("category_id", "99")
            best_matched = matched
            best_dates = extract_dates_from_text(content, date_patterns) if date_patterns else []
            best_weight = weight

    if best_score < 0.3 and is_image_file(filepath):
        best_type = "照片"
        best_category_id = "19"
        best_score = 0.5
        best_matched = []
        best_dates = []

    confidence = round(best_score, 4)
    needs_review = confidence < 0.7

    all_dates = extract_all_years_from_text(content)
    if best_dates:
        all_dates = sorted(set(all_dates + best_dates))

    if best_score < 0.15 and not is_image_file(filepath):
        best_type = "其他"
        best_category_id = "99"

    return {
        "file": str(filepath),
        "file_type": ext.lstrip('.'),
        "classification": best_type,
        "category_id": best_category_id,
        "confidence": confidence,
        "matched_keywords": best_matched,
        "extracted_dates": all_dates,
        "needs_review": needs_review,
        "content_preview": content_preview,
        "error": "",
    }


# ============================================================
# scan 子命令
# ============================================================

def cmd_scan(args):
    scan_dir = Path(args.dir)
    if not scan_dir.is_dir():
        print(f"[scan] ERROR: 目录不存在: {args.dir}")
        return False

    exclude_dirs = set()
    if args.exclude_dir:
        for d in args.exclude_dir.split(","):
            exclude_dirs.add(d.strip())

    project_name = args.project or Path(args.dir).name

    print(f"[scan] 扫描目录: {scan_dir}")
    print(f"[scan] 排除目录: {exclude_dirs}")
    print(f"[scan] 项目名: {project_name}")

    results = []
    total_files = 0
    processed = 0

    for filepath in scan_dir.rglob("*"):
        if not filepath.is_file():
            continue

        rel_parts = filepath.relative_to(scan_dir).parts
        skip = False
        for part in rel_parts:
            if part in exclude_dirs:
                skip = True
                break
        if skip:
            continue

        ext = filepath.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS and not is_image_file(filepath):
            continue

        total_files += 1

    print(f"[scan] 找到 {total_files} 个候选文件")

    for filepath in scan_dir.rglob("*"):
        if not filepath.is_file():
            continue

        rel_parts = filepath.relative_to(scan_dir).parts
        skip = False
        for part in rel_parts:
            if part in exclude_dirs:
                skip = True
                break
        if skip:
            continue

        ext = filepath.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS and not is_image_file(filepath):
            continue

        processed += 1
        rel_path = filepath.relative_to(scan_dir)
        print(f"[scan] [{processed}/{total_files}] {rel_path}")

        result = classify_single_file(filepath, project_name=project_name)
        result["relative_path"] = str(rel_path)
        results.append(result)

    class_counts = {}
    for r in results:
        cls = r["classification"]
        class_counts[cls] = class_counts.get(cls, 0) + 1

    output_data = {
        "scan_time": datetime.now().isoformat(),
        "scan_directory": str(scan_dir),
        "project_name": project_name,
        "total_files_scanned": len(results),
        "classifications": class_counts,
        "needs_review_count": 0,
        "files": results,
    }
    output_data["needs_review_count"] = sum(1 for r in results if r["needs_review"])

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(output_data, ensure_ascii=False, indent=2), encoding='utf-8')

    print(f"\n[scan] === 扫描完成 ===")
    print(f"[scan] 总文件数: {len(results)}")
    print(f"[scan] 需人工复查: {output_data['needs_review_count']}")
    for cls, count in sorted(class_counts.items()):
        print(f"[scan]   {cls}: {count}")
    print(f"[scan] 结果已保存: {output_path}")

    return True


# ============================================================
# filter 子命令
# ============================================================

def cmd_filter(args):
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[filter] ERROR: 输入文件不存在: {args.input}")
        return False

    application_year = args.year
    recent_three = [application_year - 3, application_year - 2, application_year - 1]

    print(f"[filter] 申报年度: {application_year}")
    print(f"[filter] 近三年范围: {recent_three[0]}-{recent_three[2]}")

    data = json.loads(input_path.read_text(encoding='utf-8'))
    files = data.get("files", [])

    target_dir = Path(args.target_dir)
    expired_dir = target_dir / "_过期资料"
    report_lines = []

    report_lines.append(f"# 资料整理报告")
    report_lines.append(f"")
    report_lines.append(f"- **整理时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append(f"- **申报年度**: {application_year}")
    report_lines.append(f"- **有效日期范围**: {recent_three[0]}年 - {recent_three[2]}年")
    report_lines.append(f"- **分类来源**: {args.input}")
    report_lines.append(f"- **总文件数**: {len(files)}")
    report_lines.append(f"")

    valid_files = []
    expired_files = []
    undated_files = []
    other_files = []

    for f in files:
        classification = f.get("classification", "其他")
        extracted_dates = f.get("extracted_dates", [])

        if classification in ("照片", "营业执照", "学历证书", "管理制度", "设备清单"):
            undated_files.append(f)
            continue

        if not extracted_dates:
            has_date_patterns = False
            if classification in CLASSIFICATION_PATTERNS:
                has_date_patterns = bool(CLASSIFICATION_PATTERNS[classification].get("date_patterns"))
            if has_date_patterns:
                expired_files.append(f)
            else:
                undated_files.append(f)
            continue

        max_year = max(extracted_dates)
        f["max_extracted_year"] = max_year

        if max_year < min(recent_three):
            expired_files.append(f)
        elif max_year > max(recent_three):
            undated_files.append(f)
        else:
            valid_files.append(f)

    all_valid = valid_files + undated_files

    report_lines.append(f"## 统计")
    report_lines.append(f"")
    report_lines.append(f"| 类别 | 数量 |")
    report_lines.append(f"|------|------|")
    report_lines.append(f"| 有效资料（近三年内） | {len(valid_files)} |")
    report_lines.append(f"| 无日期资料（保留） | {len(undated_files)} |")
    report_lines.append(f"| 过期资料 | {len(expired_files)} |")
    report_lines.append(f"| **合计** | **{len(all_valid)}** |")
    report_lines.append(f"")

    categorized = {}
    for f in all_valid:
        cat_id = f.get("category_id", "99")
        if cat_id not in categorized:
            categorized[cat_id] = []
        categorized[cat_id].append(f)

    report_lines.append(f"## 按类别分布")
    report_lines.append(f"")
    report_lines.append(f"| 类别 | 数量 |")
    report_lines.append(f"|------|------|")
    for cat_id in sorted(categorized.keys(), key=lambda x: int(x) if x.isdigit() else 999):
        cat_name = FILE_ORGANIZER_CATEGORIES.get(cat_id, {}).get("name", f"类别{cat_id}")
        report_lines.append(f"| {cat_name} | {len(categorized[cat_id])} |")
    report_lines.append(f"")

    if args.copy:
        print(f"[filter] 复制文件到: {target_dir}")
        copied_count = 0
        error_count = 0

        target_dir.mkdir(parents=True, exist_ok=True)

        for cat_id, cat_files in categorized.items():
            cat_name = FILE_ORGANIZER_CATEGORIES.get(cat_id, {}).get("name", f"{cat_id}.其他")
            cat_dir = target_dir / cat_name
            cat_dir.mkdir(parents=True, exist_ok=True)

            for f in cat_files:
                src = Path(f["file"])
                if not src.exists():
                    error_count += 1
                    continue
                dst = cat_dir / src.name
                counter = 1
                while dst.exists():
                    stem = src.stem
                    dst = cat_dir / f"{stem}_{counter}{src.suffix}"
                    counter += 1
                try:
                    shutil.copy2(src, dst)
                    copied_count += 1
                except Exception as e:
                    error_count += 1
                    print(f"[filter] 复制失败: {src} -> {dst}: {e}")

        if expired_files:
            expired_dir.mkdir(parents=True, exist_ok=True)
            for f in expired_files:
                src = Path(f["file"])
                if not src.exists():
                    continue
                dst = expired_dir / src.name
                counter = 1
                while dst.exists():
                    stem = src.stem
                    dst = expired_dir / f"{stem}_{counter}{src.suffix}"
                    counter += 1
                try:
                    shutil.copy2(src, dst)
                except Exception as e:
                    print(f"[filter] 复制过期文件失败: {src}: {e}")

        print(f"[filter] 复制完成: {copied_count} 个文件, 失败 {error_count} 个")
        report_lines.append(f"## 复制结果")
        report_lines.append(f"")
        report_lines.append(f"- 成功复制: {copied_count} 个文件")
        report_lines.append(f"- 失败: {error_count} 个")
        if expired_files:
            report_lines.append(f"- 过期文件归档: {len(expired_files)} 个文件 → `_过期资料/`")
        report_lines.append(f"")

    report_lines.append(f"## 需人工复查的文件")
    report_lines.append(f"")
    for f in all_valid:
        if f.get("needs_review"):
            report_lines.append(f"- `{f.get('relative_path', f.get('file', ''))}` "
                                f"({f.get('classification', '其他')}, 置信度: {f.get('confidence', 0)})")
    report_lines.append(f"")

    if expired_files:
        report_lines.append(f"## 过期文件清单")
        report_lines.append(f"")
        for f in expired_files:
            report_lines.append(f"- `{f.get('relative_path', f.get('file', ''))}` "
                                f"({f.get('classification', '其他')}, "
                                f"提取年份: {f.get('extracted_dates', [])})")
        report_lines.append(f"")

    report_path = target_dir / "_整理报告.md"
    report_path.write_text("\n".join(report_lines), encoding='utf-8')
    print(f"[filter] 报告已生成: {report_path}")

    return True


# ============================================================
# main
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='基于文件内容的资料分类器 - 高新技术企业认定',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例：
  python file_content_classifier.py scan --dir "D:\\\\资料" --output "分类结果.json" --exclude-dir ".trae,__pycache__"
  python file_content_classifier.py filter --input "分类结果.json" --year 2026 --target-dir "D:\\\\整理输出" --project-root "D:\\\\项目根目录" --copy
        ''',
    )
    subparsers = parser.add_subparsers(dest="command", help="子命令")

    scan_parser = subparsers.add_parser("scan", help="扫描目录，基于文件内容分类")
    scan_parser.add_argument("--dir", required=True, help="要扫描的目录路径")
    scan_parser.add_argument("--output", default="classification_result.json", help="输出JSON文件路径")
    scan_parser.add_argument("--exclude-dir", default=".trae", help="排除的目录，逗号分隔")
    scan_parser.add_argument("--project", default=None, help="项目名（用于OCR缓存隔离）")

    filter_parser = subparsers.add_parser("filter", help="按日期筛选并整理分类结果")
    filter_parser.add_argument("--input", required=True, help="scan命令输出的JSON文件")
    filter_parser.add_argument("--year", type=int, required=True, help="申报年度（如 2026）")
    filter_parser.add_argument("--target-dir", required=True, help="输出目标目录")
    filter_parser.add_argument("--project-root", default=None, help="项目根目录")
    filter_parser.add_argument("--copy", action="store_true", help="执行文件复制操作")

    args = parser.parse_args()

    if args.command == "scan":
        success = cmd_scan(args)
    elif args.command == "filter":
        success = cmd_filter(args)
    else:
        parser.print_help()
        success = False

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
