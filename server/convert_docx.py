import sys, json
try:
    from docx import Document
    doc = Document(sys.argv[1])
    html = '<div style="font-family:system-ui,sans-serif;font-size:14px;line-height:1.8;color:#d1d5db;padding:16px;max-width:800px;margin:0 auto">'
    def escape(s):
        return (s or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    for p in doc.paragraphs:
        t = escape(p.text)
        if not t:
            html += '<br>'
            continue
        if p.style and p.style.name and 'Heading' in p.style.name:
            lv = p.style.name.split()[-1]
            fs = {'1': '22px', '2': '18px', '3': '16px'}.get(lv, '14px')
            html += f'<h{lv} style="margin:16px 0 8px;font-size:{fs};font-weight:600;color:#f3f4f6">{t}</h{lv}>'
        else:
            html += f'<p style="margin:4px 0">{t}</p>'
    for table in doc.tables:
        html += '<table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:13px">'
        for i, row in enumerate(table.rows):
            html += '<tr>'
            for cell in row.cells:
                tag = 'th' if i == 0 else 'td'
                bg = '#1f2937' if i == 0 else 'transparent'
                html += f'<{tag} style="border:1px solid #374151;padding:6px 10px;text-align:left;background:{bg}">{escape(cell.text)}</{tag}>'
            html += '</tr>'
        html += '</table>'
    html += '</div>'
    print(json.dumps({"ok": True, "html": html}))
except Exception as e:
    print(json.dumps({"ok": False, "error": str(e)}))
