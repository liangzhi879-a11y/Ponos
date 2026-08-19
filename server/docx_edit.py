"""Word 应用内编辑：块结构读取 + 文本写回（保留段落/表格结构样式）。

用法:
  docx_edit.py read  <path>          # 输出 {"ok":true,"blocks":[{kind,text?|rows?}]}
  docx_edit.py write <jsonPath>      # 入参 {"path","blocks":[...]}

- read: 段落（Heading 1/2/3 -> h1/h2/h3，其余 p）+ 表格（cell 文本二维）
- write: 按块序更新段落 text（保留段落样式；段落内 run 级格式归一段落默认）
         与表格 cell 文本；图片/嵌入对象不触碰
"""
import sys, json


def read_docx(path):
    from docx import Document
    doc = Document(path)
    blocks = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        kind = 'p'
        if 'Heading 1' in style:
            kind = 'h1'
        elif 'Heading 2' in style:
            kind = 'h2'
        elif 'Heading 3' in style:
            kind = 'h3'
        blocks.append({"kind": kind, "text": p.text})
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        blocks.append({"kind": "table", "rows": rows})
    return {"ok": True, "blocks": blocks}


def set_para_text(p, new_text):
    """整段改写文本：写入第一个有文本的 run（按索引），其余 run 文本清空（保留段落样式）。
    纯图片/无文本 run 的段落跳过，不破坏嵌入对象。
    注意：p.runs 每次访问重建包装对象，必须按索引操作，不能比较对象同一性。"""
    runs = p.runs
    if not runs:
        return
    texts = [r.text for r in runs]
    if not any(texts):
        return  # 纯图片/空段，不写
    idx = next((i for i, t in enumerate(texts) if t), 0)
    runs[idx].text = new_text
    for i, r in enumerate(runs):
        if i != idx:
            r.text = ''


def write_docx(body):
    from docx import Document
    path = body["path"]
    doc = Document(path)
    paras = [b for b in body.get("blocks", []) if b["kind"] != "table"]
    tables = [b for b in body.get("blocks", []) if b["kind"] == "table"]
    for b, p in zip(paras, doc.paragraphs):
        if b.get("text") is not None:
            set_para_text(p, b["text"])
    for b, tbl in zip(tables, doc.tables):
        for ri, row in enumerate(b.get("rows", [])):
            if ri >= len(tbl.rows):
                break
            cells = tbl.rows[ri].cells
            for ci, val in enumerate(row):
                if ci >= len(cells):
                    break
                cell = cells[ci]
                if cell.paragraphs:
                    set_para_text(cell.paragraphs[0], str(val))
    doc.save(path)
    return {"ok": True}


def main():
    mode, arg = sys.argv[1], sys.argv[2]
    if mode == "read":
        print(json.dumps(read_docx(arg)))
    else:
        with open(arg, encoding="utf-8") as f:
            body = json.load(f)
        print(json.dumps(write_docx(body)))


try:
    main()
except Exception as e:
    print(json.dumps({"ok": False, "error": str(e)}))
